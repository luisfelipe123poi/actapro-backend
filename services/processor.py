import os
import time
import hashlib
import ffmpeg
import fitz
import pdfplumber
import io
import base64
from datetime import datetime
import assemblyai as aai
from docx import Document
from database import actas_collection, users_collection, transripciones_collection, scanners_historial_collection

def background_procesar_asamblea(
    session_id: str,
    temp_audio_path: str,
    email: str,
    instrucciones: str,
    nombre_personalizado: str,
    duracion_segundos: float,
    horas_usadas_mes: float,
    content_bytes: bytes,
    filename: str,
    openai_client,
    PROMPT_SISTEMA_ACTAS
):
    """Tarea en segundo plano para procesar toda la asamblea sin omitir ninguna validación ni lógica original."""
    try:
        # 1. Registrar estado inicial "transcribiendo" en la base de datos para seguimiento (Polling)
        actas_collection.insert_one({
            "session_id": session_id,
            "email": email,
            "estado": "transcribiendo",
            "mensaje": "Iniciando transcripción con AssemblyAI...",
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M")
        })

        # Caché de transcripción para evitar gastos duplicados en AssemblyAI
        file_hash = hashlib.sha256(content_bytes).hexdigest()
        cached_transcription = transripciones_collection.find_one(
            {"file_hash": file_hash}
        )

        if cached_transcription:
            print("💡 Audio duplicado detectado: Reutilizando transcripción guardada para evitar gasto en AssemblyAI.")
            texto_transcrito = cached_transcription["texto_transcrito"]
        else:
            config = aai.TranscriptionConfig(speaker_labels=True, language_code="es")
            transcriber = aai.Transcriber()
            transcript = transcriber.transcribe(temp_audio_path, config=config)

            if transcript.status == aai.TranscriptStatus.error:
                raise Exception(f"Error en AssemblyAI: {transcript.error}")

            texto_transcrito = ""
            if transcript.utterances:
                for utterance in transcript.utterances:
                    texto_transcrito += f"[Persona {utterance.speaker}]: {utterance.text}\n"
            else:
                texto_transcrito = transcript.text

            transripciones_collection.insert_one({
                "file_hash": file_hash,
                "filename": filename,
                "texto_transcrito": texto_transcrito,
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "createdAt": datetime.utcnow(),
            })

        # Actualizar estado a analizando con IA
        actas_collection.update_one(
            {"session_id": session_id},
            {"$set": {"estado": "analizando_ia", "mensaje": "Estructurando minuta con Inteligencia Artificial..."}}
        )

        # Lógica de nombre del acta
        if not nombre_personalizado or nombre_personalizado.strip() == "":
            try:
                prompt_nombre = f"""
                Analiza el siguiente fragmento de transcripción de una asamblea y extrae estrictamente el nombre del edificio, conjunto residencial, copropiedad o empresa mencionada. 
                Responde ÚNICAMENTE con un nombre limpio apto para archivo (sin espacios, usa guiones bajos _, sin tildes ni caracteres especiales, por ejemplo: Acta_Asamblea_Edificio_Torre_Central).
                
                Transcripción:
                {texto_transcrito[:2500]}...
                """
                resp_nombre = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt_nombre}],
                    temperature=0.2,
                )
                nombre_ia = resp_nombre.choices[0].message.content.strip().replace(" ", "_")
                nombre_ia = "".join(c for c in nombre_ia if c.isalnum() or c in ("_", "-"))
                nombre_archivo_acta = f"{nombre_ia}.docx" if nombre_ia else f"Acta_Asamblea_{session_id[:8]}.docx"
            except Exception:
                nombre_archivo_acta = f"Acta_Asamblea_{session_id[:8]}.docx"
        else:
            nombre_limpio = nombre_personalizado.strip().replace(" ", "_")
            nombre_limpio = "".join(c for c in nombre_limpio if c.isalnum() or c in ("_", "-", "."))
            nombre_base = nombre_limpio.replace(".docx", "")
            nombre_archivo_acta = f"{nombre_base}.docx"

        output_docx_path = f"temp_outputs/{session_id}_{nombre_archivo_acta}"

        prompt_sistema = PROMPT_SISTEMA_ACTAS
        if instrucciones:
            prompt_sistema += f"\n\nINSTRUCCIONES ADICIONALES DEL USUARIO:\n{instrucciones}"

        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": f"Transcripción de la asamblea:\n\n{texto_transcrito}"},
            ],
            temperature=0.3,
        )

        acta_final = response.choices[0].message.content

        # Actualizar estado a generando documento
        actas_collection.update_one(
            {"session_id": session_id},
            {"$set": {"estado": "generando_docx", "mensaje": "Generando documento Word..."}}
        )

        doc = Document()
        titulo_principal = doc.add_heading("ACTA DE ASAMBLEA GENERAL DE COPROPIETARIOS", level=0)
        titulo_principal.alignment = 1

        for linea in acta_final.split("\n"):
            linea_clean = linea.strip()
            if not linea_clean:
                continue

            if linea_clean.startswith("# "):
                doc.add_heading(linea_clean.replace("# ", "").strip(), level=1)
            elif linea_clean.startswith("## ") or linea_clean.startswith("### "):
                doc.add_heading(linea_clean.replace("#", "").strip(), level=2)
            else:
                p = doc.add_paragraph()
                if "**" in linea_clean:
                    partes = linea_clean.split("**")
                    for i, parte in enumerate(partes):
                        if parte:
                            run = p.add_run(parte)
                            if i % 2 == 1:
                                run.bold = True
                else:
                    p.add_run(linea_clean)

        doc.save(output_docx_path)
        peso_archivo = f"{round(os.path.getsize(output_docx_path) / 1024, 1)} KB"

        # Guardar el documento completo en la colección de actas
        actas_collection.update_one(
            {"session_id": session_id},
            {
                "$set": {
                    "estado": "completado",
                    "mensaje": "¡Acta generada con éxito!",
                    "nombre_acta": nombre_archivo_acta,
                    "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "peso": peso_archivo,
                    "contenido": acta_final,
                    "output_path": output_docx_path
                }
            }
        )

        # 4. Actualizar consumo mensual de horas sumando la duración real del audio procesado
        nuevas_horas = horas_usadas_mes + (duracion_segundos / 3600.0)
        users_collection.update_one(
            {"email": email}, 
            {"$set": {"horas_usadas_mes": nuevas_horas}}
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        actas_collection.update_one(
            {"session_id": session_id},
            {"$set": {"estado": "error", "mensaje": str(e)}}
        )
    finally:
        if os.path.exists(temp_audio_path):
            os.remove(temp_audio_path)


def background_escanear_documento(
    scanner_id: str,
    file_bytes: bytes,
    filename: str,
    email: str,
    client_openai
):
    """Tarea en segundo plano para procesar el escaneo sin omitir ningún prompt ni regla original."""
    try:
        texto_extraido = ""
        es_imagen = filename.endswith((".png", ".jpg", ".jpeg", ".webp"))
        
        # 2. Extracción según el formato del archivo (Idéntico al original)
        if es_imagen:
            base64_image = base64.b64encode(file_bytes).decode("utf-8")
            contenido_usuario = [
                {
                    "type": "text",
                    "text": "Analyze this scanned document or image. Extract the information by structuring the visual design with corporate semantic HTML tags (use <h1>, <h2>, <p>, <table>, <thead>, <tbody>, <tr>, <th>, <td>). Apply Tailwind CSS classes to maintain a professional style (e.g., fonts, clean borders, spacing). DO NOT use Markdown, DO NOT use asterisks, DO NOT use code blocks of any kind. If there are data tables, create them completely with HTML tags. If you detect statistical charts or diagrams, represent them with a structured div with the class 'p-4 border-2 border-dashed border-slate-300 bg-slate-50 text-center text-slate-500 rounded-lg text-xs my-4' indicating the content of the chart."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}"
                    }
                }
            ]
        else:
            if filename.endswith(".pdf"):
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    texto_pagina = page.get_text()
                    
                    if texto_pagina.strip():
                        texto_extraido += f"\n--- Página {page_num + 1} ---\n" + texto_pagina
                
                doc.close()
                
                # Respaldo con pdfplumber para extracción precisa de tablas en PDFs
                if len(texto_extraido.strip()) < 50:
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for i, page in enumerate(pdf.pages):
                            t = page.extract_text()
                            if t:
                                texto_extraido += f"\n--- Página (Tablas) {i + 1} ---\n" + t

            elif filename.endswith((".txt", ".doc", ".docx")):
                texto_extraido = file_bytes.decode("utf-8", errors="ignore")
            else:
                raise Exception("Formato de archivo no soportado. Sube un PDF, imagen o documento de texto.")

            if not texto_extraido.strip():
                raise Exception("El documento está vacío o no se pudo extraer texto legible.")

            contenido_usuario = f"""Analyze the following text extracted from the document. Your output must be EXCLUSIVELY corporate semantic HTML ready to render directly in a browser or web container.
- Replicate the original visual and section structure.
- Use <h1>, <h2> for main and section titles.
- Use <p> for paragraphs with Tailwind classes (e.g., text-slate-900, text-xs, leading-relaxed).
- Use complete table tags (<table>, <thead>, <tbody>, <tr>, <td>) with borders and corporate classes if there is structured data.
- If you detect references to charts, schemes, or diagrams, create them as a visual block with a dotted border.
- FORBIDDEN to use Markdown, asterisks (*), markdown list hyphens (#), or wrap the result in markdown code quotes.

Extracted text:
{texto_extraido[:15000]}"""

        # 3. Procesamiento inteligente y estructurado con OpenAI GPT-4o (Con tus prompts exactos)
        response_openai = client_openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You are a Documentation Engineer and Web Designer expert in corporate digitalization. 
Your sole mission is to transform the input document information into a pure, clean, and professional HTML code block integrated with Tailwind CSS classes.
STRICT RULES:
1. Return ONLY valid HTML code. Do not include prior explanations or text outside of the HTML.
2. NEVER use Markdown syntax (*, #, -, ```html). The result must be plain text containing exclusively HTML markup.
3. Structure data tables with <table>, <thead>, <tbody>, <tr>, <th>, and <td> applying clean classes (e.g., border border-slate-300 p-2).
4. Represent detected charts using a <div> container with dotted borders and professional design.
5. Maintain absolute fidelity to the original document structure."""
                },
                {
                    "role": "user",
                    "content": contenido_usuario
                }
            ],
            temperature=0.0
        )

        resultado_html = response_openai.choices[0].message.content.strip()

        # Descontar / sumar tokens consumidos en la BD si el usuario está autenticado
        tokens_consumidos = 0
        if email and hasattr(response_openai, "usage") and response_openai.usage:
            tokens_consumidos = response_openai.usage.total_tokens
            users_collection.update_one(
                {"email": email},
                {"$inc": {"tokens_usados": tokens_consumidos}}
            )

        # Limpieza defensiva por si el modelo por inercia agrega bloques de código
        if resultado_html.startswith("```html"):
            resultado_html = resultado_html[7:]
        if resultado_html.startswith("```"):
            resultado_html = resultado_html[3:]
        if resultado_html.endswith("```"):
            resultado_html = resultado_html[:-3]
        resultado_html = resultado_html.strip()

        # Actualizar el registro en la colección de scanners con el resultado completado
        scanners_historial_collection.update_one(
            {"scanner_id": scanner_id},
            {
                "$set": {
                    "estado": "completado",
                    "contenido": resultado_html,
                    "tokens": tokens_consumidos
                }
            }
        )

    except Exception as e:
        # Si ocurre un error en segundo plano, se guarda el estado de error
        scanners_historial_collection.update_one(
            {"scanner_id": scanner_id},
            {"$set": {"estado": "error", "mensaje": str(e)}}
        )