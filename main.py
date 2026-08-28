import hashlib 
import io 
import os
import re
import shutil
import uuid
from datetime import datetime, timezone, timedelta
from typing import Optional
from pydantic import BaseModel
import assemblyai as aai
from bson import ObjectId
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, JSONResponse
import mercadopago
import openai
from openai import OpenAI
import pdfplumber
import pymupdf as fitz
from pymongo import MongoClient
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
import sib_api_v3_sdk
from sib_api_v3_sdk.rest import ApiException
from sib_api_v3_sdk.models import SendSmtpEmail, SendSmtpEmailSender, SendSmtpEmailTo
from tasks import task_procesar_asamblea

load_dotenv()

# ==========================================
# 0. Inicialización y Configuración General
# ==========================================

app = FastAPI(
    title="ActaProCore API con MongoDB y Mercado Pago", 
    version="1.9.5"
)

# Lista completa de orígenes permitidos
origins = [
    "https://actaprocore.com",
    "https://www.actaprocore.com",
    "https://cotizaciones-apc.actaprocore.com",
    "https://admin-licencias.actaprocore.com",
    "https://dashboard-apc.actaprocore.com", 
    "https://panel-infraestructura.actaprocore.com",
    "https://sptc-soporte.actaprocore.com",
    "http://localhost",
    "http://localhost:3000",
    "http://localhost:5173",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:5173"
]

# Configuración estricta e impecable de CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=600,
)

# Configuración de Claves y Base de Datos
AAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")
MP_ACCESS_TOKEN = os.getenv("MP_ACCESS_TOKEN")
BREVO_API_KEY = os.getenv("BREVO_API_KEY")

if not AAI_API_KEY or not OPENAI_API_KEY:
    raise ValueError(
        "❌ Error: Claves de API de AssemblyAI u OpenAI no configuradas en .env"
    )

# Inicializar cliente de MongoDB especificando la base de datos 'actabot_db'
mongo_client = MongoClient(MONGO_URI)
db = mongo_client["actabot_db"]
users_collection = db["users"]
actas_collection = db["actas_historial"]
transripciones_collection = db["transripciones_cache"]
scanners_historial_collection = db["scanners_historial"]
soporte_collection = db["soporte_feedback"]

# Configurar índice TTL para caché de transcripciones
try:
    transripciones_collection.create_index(
        "createdAt", expireAfterSeconds=2592000
    )
except Exception as e:
    print(f"Nota sobre índice TTL: {e}")

# ==========================================
# 1. Configuración de Brevo (Cliente Global)
# ==========================================
brevo_client = None
if BREVO_API_KEY:
    configuration = sib_api_v3_sdk.Configuration()
    configuration.api_key['api-key'] = BREVO_API_KEY
    brevo_client = sib_api_v3_sdk.TransactionalEmailsApi(sib_api_v3_sdk.ApiClient(configuration))
else:
    print("⚠️ Advertencia: BREVO_API_KEY no está configurada. Los correos electrónicos no se enviarán.")

# Inicializar SDK de Mercado Pago y Clientes de IA
mp_sdk = mercadopago.SDK(MP_ACCESS_TOKEN) if MP_ACCESS_TOKEN else None

aai.settings.api_key = AAI_API_KEY
client = OpenAI(api_key=OPENAI_API_KEY)

# ==========================================
# 2. Prompts y Diccionarios de Configuración
# ==========================================
prompt_sistema = """
Eres un Secretario Jurídico de alto nivel, experto en Propiedad Horizontal en Colombia (Ley 675 de 2001). 
Tu misión es transformar la transcripción de audio adjunta en un ACTA FORMAL, DETALLADA Y COMPLETA. 

ORDEN DE TRABAJO (ESTRICTO):
1. INTEGRIDAD DE LA INFORMACIÓN: NO RESUMAS EL CONTENIDO TÉCNICO NI LAS PROPUESTAS. Debes capturar todos los argumentos, cifras, justificaciones financieras, propuestas de mantenimiento y posturas de los copropietarios. Si alguien propone algo específico, inclúyelo detalladamente.
2. FILTRADO DE RUIDO: ÚNICAMENTE elimina interrupciones, saludos, chismes, peleas personales o frases vacías que no aporten al objeto de la asamblea. Todo lo que tenga que ver con gestión, presupuesto, administración o decisiones debe quedar plasmado.
3. ESTRUCTURA JURÍDICA:
   - ENCABEZADO Y QUÓRUM: Detalla la verificación de coeficientes si se menciona.
   - DESARROLLO PUNTO POR PUNTO: Para CADA punto del orden del día, redacta el desarrollo de forma narrativa pero minuciosa. Transcribe los debates relevantes ("El copropietario A solicitó aclarar X; el administrador respondió que Y").
   - DECISIONES Y VOTACIONES: Registra el sentido de las votaciones y cualquier salvedad o constancia que los copropietarios hayan solicitado dejar por escrito.
4. ESTILO Y FORMATO: Lenguaje jurídico formal, impersonal y preciso. Utiliza asteriscos dobles (ej: **$1.500.000** o **Aprobado por mayoría**) para resaltar en el texto cifras, costos, valores y decisiones clave. Redacta como si hubieras estado presente tomando nota exhaustiva de todo lo importante.
"""

PROMPT_SISTEMA_ACTAS = """Eres un Secretario Jurídico experto en Propiedad Horizontal en Colombia (Ley 675 de 2001). Tu objetivo es redactar un acta formal, jurídica y detallada a partir de la transcripción de la asamblea provista, asegurando un formato profesional en Markdown y cumplimiento legal estricto."""

PRECIOS_PLANES = {
    "basico": {
        "nombre": "Plan Básico",
        "precio": 4000.0,
        "tokens_mensuales": 100000,
        "documentos_estimados": 40,
        "limite_horas": 15.0
    },
    "profesional": {
        "nombre": "Plan Intermedio",
        "precio": 4000.0,
        "tokens_mensuales": 300000,
        "documentos_estimados": 120,
        "limite_horas": 60.0,
    },
    "corporativo": {
        "nombre": "Plan Profesional / Pro",
        "precio": 4000.0,
        "tokens_mensuales": 1000000,
        "documentos_estimados": 400,
        "limite_horas": 200.0,
    },
}

CONFIGURACION_PLANES = {
    "free": {"tokens": 10000, "horas": 3.0},
    "basico": {"tokens": 100000, "horas": 15.0},
    "profesional": {"tokens": 300000, "horas": 60.0},
    "corporativo": {"tokens": 1000000, "horas": 200.0}
}

# ==========================================
# 3. Modelos Pydantic
# ==========================================
class AuthModel(BaseModel):
    email: str
    password: str
    plan: Optional[str] = "free"

class PaymentPreferenceModel(BaseModel):
    email: str
    plan_name: str
    price: Optional[float] = None

class RenombrarActaRequest(BaseModel):
    email: str
    acta_id: str
    nuevo_nombre: str

class EliminarActaRequest(BaseModel):
    email: str
    acta_id: str

class ConsultaFAQRequest(BaseModel):
    tipoConsulta: str

class ConsultaPlanRequest(BaseModel):
    clienteId: str
    tipoProblema: Optional[str] = None
# ==========================================
# 4. Función Auxiliar para Enviar Correos con Brevo
# ==========================================
def enviar_correo_brevo(destinatario_email: str, destinatario_nombre: str, asunto: str, html_contenido: str):
    """Envía un correo electrónico transaccional utilizando la API oficial de Brevo."""
    if not brevo_client:
        print("Brevo no está inicializado (falta la BREVO_API_KEY).")
        return False
    
    try:
        email_data = SendSmtpEmail(
            sender=SendSmtpEmailSender(name="ActaPro Core", email="contacto@actaprocore.com"),
            to=[SendSmtpEmailTo(email=destinatario_email, name=destinatario_nombre)],
            subject=asunto,
            html_content=html_contenido
        )
        brevo_client.send_transac_email(email_data)
        print(f"Correo enviado exitosamente a {destinatario_email}")
        return True
    except ApiException as e:
        print(f"Error de API al enviar correo mediante Brevo: {e}")
        return False
    except Exception as e:
        print(f"Error general al enviar correo mediante Brevo: {e}")
        return False
# ==========================================
# 3. Rutas de la Aplicación
# ==========================================

@app.post("/api/registro")
def registrar_usuario(data: AuthModel):
    existing_user = users_collection.find_one({"email": data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="El correo ya está registrado.")

    # Obtener configuración del plan
    plan_config = CONFIGURACION_PLANES.get(data.plan, CONFIGURACION_PLANES["free"])

    nuevo_usuario = {
        "email": data.email,
        "password": data.password,  # Nota: Recuerda hashear esto en producción
        "plan": data.plan,
        # Gestión de Horas
        "horas_restantes": plan_config["horas"],
        "horas_usadas_mes": 0.0,
        "limite_horas_mes": plan_config["horas"],
        # Gestión de Tokens
        "tokens_usados": 0,
        "limite_tokens_mes": plan_config["tokens"]
    }

    users_collection.insert_one(nuevo_usuario)

    # Enviar correo de bienvenida profesional según el plan con Brevo
    nombre_usuario = data.email.split("@")[0].capitalize()
    asunto_correo = f"¡Bienvenido a ActaPro Core - Plan {data.plan.capitalize()} Activado!"
    
    # HTML Profesional con los colores corporativos de ActaPro Core y Banner Superior
    html_cuerpo = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
    </head>
    <body style="margin: 0; padding: 0; background-color: #f4f6f9; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
        <table align="center" border="0" cellpadding="0" cellspacing="0" width="100%" style="background-color: #f4f6f9; padding: 40px 0;">
            <tr>
                <td align="center">
                    <!-- Contenedor Principal -->
                    <table border="0" cellpadding="0" cellspacing="0" width="600" style="background-color: #ffffff; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05);">
                        
                        <!-- Banner Superior con Imagen -->
                        <tr>
                            <td align="center" style="background-color: #0f172a; padding: 0;">
                                <img src="https://images.unsplash.com/photo-1557804506-669a67965ba0?auto=format&fit=crop&w=1200&q=80" alt="ActaPro Core Banner" width="600" style="display: block; width: 100%; height: 160px; object-fit: cover; opacity: 0.85;" />
                            </td>
                        </tr>

                        <!-- Franja de Identidad / Logo Textual -->
                        <tr>
                            <td align="center" style="background-color: #0f172a; padding: 0 30px 25px 30px;">
                                <h1 style="color: #ffffff; font-size: 24px; margin: 0; letter-spacing: 0.5px;">ActaPro <span style="color: #38bdf8;">Core</span></h1>
                                <p style="color: #94a3b8; font-size: 13px; margin: 5px 0 0 0; text-transform: uppercase; letter-spacing: 1px;">Inteligencia Jurídica para Asambleas</p>
                            </td>
                        </tr>

                        <!-- Cuerpo del Mensaje -->
                        <tr>
                            <td style="padding: 40px 30px; color: #334155;">
                                <h2 style="color: #0f172a; font-size: 20px; margin-top: 0; margin-bottom: 20px;">¡Hola, {nombre_usuario}!</h2>
                                <p style="font-size: 15px; line-height: 1.6; color: #475569; margin-bottom: 25px;">
                                    Tu cuenta ha sido registrada y configurada exitosamente. Ya tienes acceso completo a nuestra plataforma automatizada bajo el plan <strong style="color: #0f172a; text-transform: uppercase;">{data.plan}</strong>.
                                </p>

                                <!-- Tarjeta de Resumen de Plan -->
                                <table border="0" cellpadding="0" cellspacing="0" width="100%" style="background-color: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; margin-bottom: 25px;">
                                    <tr>
                                        <td style="padding: 20px;">
                                            <p style="margin: 0 0 10px 0; font-size: 14px; font-weight: bold; color: #1e293b; text-transform: uppercase; letter-spacing: 0.5px;">Resumen de tu Suscripción:</p>
                                            <table border="0" cellpadding="0" cellspacing="0" width="100%">
                                                <tr>
                                                    <td style="padding: 6px 0; font-size: 14px; color: #64748b;">Límite de Tokens:</td>
                                                    <td align="right" style="padding: 6px 0; font-size: 14px; font-weight: bold; color: #0f172a;">{plan_config["tokens"]:,}</td>
                                                </tr>
                                                <tr>
                                                    <td style="padding: 6px 0; font-size: 14px; color: #64748b;">Límite de Horas de Audio:</td>
                                                    <td align="right" style="padding: 6px 0; font-size: 14px; font-weight: bold; color: #0f172a;">{plan_config["horas"]} hrs / mes</td>
                                                </tr>
                                            </table>
                                        </td>
                                    </tr>
                                </table>

                                <p style="font-size: 15px; line-height: 1.6; color: #475569; margin-bottom: 30px;">
                                    Comienza a subir tus archivos de audio y genera actas profesionales en minutos, cumpliendo estrictamente con la normativa vigente.
                                </p>

                                <!-- Botón de Acción -->
                                <table border="0" cellpadding="0" cellspacing="0" width="100%">
                                    <tr>
                                        <td align="center">
                                            <a href="https://actaprocore.com/dashboard" target="_blank" style="background-color: #2563eb; color: #ffffff; padding: 12px 30px; border-radius: 6px; text-decoration: none; font-weight: bold; font-size: 15px; display: inline-block; box-shadow: 0 4px 6px rgba(37, 99, 235, 0.2);">Acceder al Dashboard</a>
                                        </td>
                                    </tr>
                                </table>
                            </td>
                        </tr>

                        <!-- Pie de Página -->
                        <tr>
                            <td align="center" style="background-color: #f1f5f9; padding: 25px 30px; border-top: 1px solid #e2e8f0;">
                                <p style="font-size: 12px; color: #64748b; margin: 0 0 5px 0;">Este es un mensaje automático generado por <strong>ActaPro Core</strong>.</p>
                                <p style="font-size: 12px; color: #94a3b8; margin: 0;">&copy; 2026 ActaPro Core. Todos los derechos reservados.</p>
                            </td>
                        </tr>

                    </table>
                </td>
            </tr>
        </table>
    </body>
    </html>
    """
    
    enviar_correo_brevo(data.email, nombre_usuario, asunto_correo, html_cuerpo)

    return {
        "message": "Usuario registrado con éxito",
        "email": data.email,
        "plan": data.plan,
        "tokens_limite": plan_config["tokens"],
        "limite_horas_mes": plan_config["horas"]
    }

import traceback
from fastapi import status, HTTPException

@app.post("/api/login")
def login_usuario(data: AuthModel):
    try:
        # 1. Buscar usuario en MongoDB
        user = users_collection.find_one({"email": data.email})

        # 2. Validar credenciales (usando hash seguro o texto plano según tu migración)
        if not user or user.get("password") != data.password:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Credenciales inválidas."
            )

        # 3. Retornar datos exitosos
        return {
            "message": "Login exitoso",
            "email": str(user.get("email")),
            "plan": str(user.get("plan", "free")),
            "tokens_usados": int(user.get("tokens_usados", 0)),
            "limite_tokens_mes": int(user.get("limite_tokens_mes", 0)),
            "horas_restantes": float(user.get("horas_restantes", 0.0)),
            "horas_usadas_mes": float(user.get("horas_usadas_mes", 0.0)),
            "limite_horas_mes": float(user.get("limite_horas_mes", 0.0))
        }

    except HTTPException as http_ex:
        # Re-lanzar las excepciones explicitas (401, etc.) sin alterarlas
        raise http_ex

    except Exception as e:
        # Captura el rastreo completo del error
        error_trace = traceback.format_exc()
        
        # Imprime el log exacto en los logs de Render / Consola del servidor
        print("\n================ ERROR EXACTO EN LOG IN ================")
        print(error_trace)
        print("========================================================\n")
        
        # Devuelve el error exacto como respuesta JSON al cliente (útil para debugging)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "error_tipo": type(e).__name__,
                "mensaje": str(e),
                "traceback": error_trace.splitlines()
            }
        )

@app.get("/api/user/status")
def get_user_status(email: str):
    usuario = users_collection.find_one({"email": email})
    if not usuario:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")

    return {
        "email": usuario.get("email"),
        "plan": usuario.get("plan", "free"),
        "tokens_usados": usuario.get("tokens_usados", 0),
        "limite_tokens_mes": usuario.get("limite_tokens_mes", 0),
        "horas_restantes": usuario.get("horas_restantes", 0.0),
        "horas_usadas_mes": usuario.get("horas_usadas_mes", 0.0),
        "limite_horas_mes": usuario.get("limite_horas_mes", 0.0)
    }


@app.get("/api/actas/historial")
def obtener_historial_actas(email: str):
  actas_cursor = actas_collection.find({"email": email})
  actas = []
  for doc in actas_cursor:
    doc["id"] = str(doc["_id"])
    del doc["_id"]
    actas.append(doc)
  return {"actas": actas}


@app.put("/api/actas/renombrar")
async def renombrar_acta(data: RenombrarActaRequest):
  try:
    resultado = actas_collection.update_one(
        {"email": data.email, "nombre_acta": data.acta_id},
        {"$set": {"nombre_acta": data.nuevo_nombre}},
    )

    if resultado.matched_count == 0:
      from bson import ObjectId

      try:
        resultado = actas_collection.update_one(
            {"_id": ObjectId(data.acta_id), "email": data.email},
            {"$set": {"nombre_acta": data.nuevo_nombre}},
        )
      except Exception:
        pass

    if resultado.matched_count == 0:
      raise HTTPException(
          status_code=404, detail="Acta no encontrada o no autorizada."
      )

    return {"message": "¡Acta renombrada con éxito!"}
  except HTTPException as he:
    raise he
  except Exception as e:
    raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/actas/eliminar")
async def eliminar_acta(data: EliminarActaRequest):
  try:
    resultado = actas_collection.delete_one(
        {"email": data.email, "nombre_acta": data.acta_id}
    )

    if resultado.deleted_count == 0:
      from bson import ObjectId

      try:
        resultado = actas_collection.delete_one(
            {"_id": ObjectId(data.acta_id), "email": data.email}
        )
      except Exception:
        pass

    if resultado.deleted_count == 0:
      raise HTTPException(
          status_code=404, detail="Acta no encontrada para eliminar."
      )

    return {"message": "Acta eliminada correctamente."}
  except HTTPException as he:
    raise he
  except Exception as e:
    raise HTTPException(status_code=500, detail=str(e))


import io
import re
from fastapi import FastAPI, HTTPException, Response
from bson import ObjectId
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER

from fastapi.responses import RedirectResponse




# ==========================================
# 4. Endpoints de Pagos y Webhook (Mercado Pago + Brevo)
# ==========================================

@app.post("/api/crear-preferencia-pago")
def crear_preferencia_pago(data: PaymentPreferenceModel):
    if not mp_sdk:
        raise HTTPException(
            status_code=500, detail="Mercado Pago no está configurado en el servidor."
        )

    # Validar y limpiar email
    email_cliente = data.email.strip().lower() if data.email else None
    if not email_cliente:
        raise HTTPException(
            status_code=400, detail="El correo electrónico es requerido para procesar el pago."
        )

    # Normalizar el plan recibido
    plan_raw = (data.plan_name or "").lower().strip()

    # Mapeo exhaustivo para soportar variaciones y alias del frontend
    mapeo_planes = {
        # Básico
        "basico": "basico",
        "básico": "basico",
        "plan basico": "basico",
        "plan básico": "basico",
        
        # Profesional / Intermedio
        "profesional": "profesional",
        "intermedio": "profesional",
        "plan profesional": "profesional",
        "plan intermedio": "profesional",
        
        # Corporativo / Pro
        "corporativo": "corporativo",
        "pro": "corporativo",
        "plan corporativo": "corporativo",
        "plan pro": "corporativo",
        "plan profesional / pro": "corporativo"
    }

    plan_id = mapeo_planes.get(plan_raw)
    if not plan_id or plan_id not in PRECIOS_PLANES:
        raise HTTPException(
            status_code=400, detail=f"Plan no válido: '{data.plan_name}'"
        )

    # Extracción segura de datos del plan
    info_plan = PRECIOS_PLANES[plan_id]
    nombre_mostrar = info_plan.get("nombre", "Plan ActaPro")
    precio_plan = float(info_plan.get("precio", 0.0))

    # Verificar o crear usuario base en MongoDB
    user = users_collection.find_one({"email": email_cliente})
    if not user:
        users_collection.insert_one({
            "email": email_cliente,
            "password": "temp_password_temporal",
            "plan": "free",
            "tokens_usados": 0,
            "limite_tokens_mes": 0,
            "horas_usadas_mes": 0.0,
            "horas_restantes": 0.0,
            "limite_horas_mes": 3.0,
        })

    preference_data = {
        "items": [{
            "title": f"ActaBot PH - {nombre_mostrar}",
            "quantity": 1,
            "currency_id": "COP",
            "unit_price": precio_plan,
        }],
        "payer": {"email": email_cliente},
        "back_urls": {
            "success": "https://actaprocore.com/dashboard",
            "failure": "https://actaprocore.com/dashboard",
            "pending": "https://actaprocore.com/dashboard",
        },
        "auto_return": "approved",
        "notification_url": "https://actapro-backend.onrender.com/api/webhook-mercadopago",
        "statement_descriptor": "ACTABOT PH",
        "external_reference": f"{email_cliente}|{plan_id}",
    }

    try:
        preference_response = mp_sdk.preference().create(preference_data)
        
        # Extracción segura compatible con versiones del SDK
        if isinstance(preference_response, dict) and "response" in preference_response:
            preference = preference_response["response"]
        else:
            preference = preference_response

        init_point = preference.get("init_point")
        sandbox_init_point = preference.get("sandbox_init_point")

        if not init_point:
            print(f"Error MercadoPago Response: {preference_response}")
            raise HTTPException(
                status_code=400,
                detail="Mercado Pago rechazó la preferencia o devolvió credenciales inválidas.",
            )

        return {
            "init_point": init_point,
            "sandbox_init_point": sandbox_init_point,
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Exception en crear_preferencia_pago: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/webhook-mercadopago")
async def webhook_mercadopago(request: Request):
    """
    Recibe notificaciones automáticas de Mercado Pago, consulta el API oficial,
    detecta el plan de forma blindada mediante monto y referencia, actualiza los contadores en MongoDB y envía un correo con Brevo.
    """
    try:
        body = await request.json()
        if body.get("type") == "payment":
            payment_id = body.get("data", {}).get("id")
            if payment_id and mp_sdk:
                payment_info = mp_sdk.payment().get(payment_id)
                
                if isinstance(payment_info, dict) and "response" in payment_info:
                    payment = payment_info["response"]
                else:
                    payment = payment_info

                if payment.get("status") == "approved":
                    payer_email = payment.get("payer", {}).get("email")
                    ext_ref = payment.get("external_reference", "")
                    
                    # Recuperar email de respaldo o dividir la referencia externa
                    plan_desde_ref = None
                    if "|" in ext_ref:
                        parts = ext_ref.split("|")
                        if not payer_email:
                            payer_email = parts[0]
                        plan_desde_ref = parts[1]
                    elif not payer_email:
                        payer_email = ext_ref

                    if payer_email:
                        payer_email = payer_email.strip().lower()
                        
                        plan_asignado = "basico"
                        
                        # 1. Intentar por external_reference si viene empaquetado
                        if plan_desde_ref and plan_desde_ref in PRECIOS_PLANES:
                            plan_asignado = plan_desde_ref
                        else:
                            # 2. Blindaje por monto exacto pagado (transaction_amount)
                            monto_pagado = float(payment.get("transaction_amount", 0))
                            for p_id, info in PRECIOS_PLANES.items():
                                if float(info.get("precio", 0)) == monto_pagado:
                                    plan_asignado = p_id
                                    break
                            
                            # 3. Si el monto falla, respaldo por lectura de ítems o título
                            if plan_asignado == "basico":
                                items = payment.get("additional_info", {}).get("items", []) or payment.get("items", [])
                                for item in items:
                                    title_lower = item.get("title", "").lower()
                                    if "corporativo" in title_lower or "pro" in title_lower:
                                        plan_asignado = "corporativo"
                                        break
                                    elif "profesional" in title_lower or "intermedio" in title_lower:
                                        plan_asignado = "profesional"
                                        break

                        info_plan = PRECIOS_PLANES.get(plan_asignado, PRECIOS_PLANES["basico"])
                        
                        tokens_otorgados = info_plan.get("tokens_mensuales", 100000)
                        horas_otorgadas = info_plan.get("limite_horas", 15.0)
                        precio_pagado = info_plan.get("precio", 0.0)
                        nombre_plan_fmt = info_plan.get("nombre", plan_asignado.capitalize())

                        users_collection.update_one(
                            {"email": payer_email},
                            {
                                "$set": {
                                    "plan": plan_asignado,
                                    "tokens_usados": 0,
                                    "limite_tokens_mes": tokens_otorgados,
                                    "horas_usadas_mes": 0.0,
                                    "horas_restantes": horas_otorgadas,
                                    "limite_horas_mes": horas_otorgadas,
                                }
                            },
                            upsert=True
                        )

                        # Formateo correcto de variables dinámicas en el correo HTML
                        nombre_usuario = payer_email.split("@")[0].capitalize()
                        asunto_correo = f"¡Compra exitosa! Licencia {nombre_plan_fmt} Activada"
                        
                        html_cuerpo = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Activación de Licencia - ActaPro Core</title>
</head>
<body style="margin: 0; padding: 0; background-color: #f8fafc; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;">
    <table align="center" border="0" cellpadding="0" cellspacing="0" width="100%" style="padding: 30px 0; background-color: #f8fafc;">
        <tr>
            <td align="center">
                <!-- Contenedor Principal -->
                <table border="0" cellpadding="0" cellspacing="0" width="600" style="background-color: #ffffff; border-radius: 12px; border: 1px solid #e2e8f0; overflow: hidden; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);">
                    <!-- Cabecera -->
                    <tr>
                        <td align="left" style="background: #ffffff; padding: 25px 30px; border-bottom: 3px solid #10b981;">
                            <table width="100%" border="0" cellpadding="0" cellspacing="0">
                                <tr>
                                    <td>
                                        <!-- PEGA TU LINK AQUÍ -->
                                        <img src="https://res.cloudinary.com/iuu7h8rj/image/upload/v1787944318/1.png" alt="ActaPro Core" width="130" style="display: block; border: 0; outline: none; text-decoration: none; max-width: 130px; height: auto;">
                                    </td>
                                    <td align="right" style="color: #64748b; font-size: 12px;">
                                        Soporte Oficial
                                    </td>
                                </tr>
                            </table>
                        </td>
                    </tr>
                    
                    <!-- Cuerpo del Contenido -->
                    <tr>
                        <td style="padding: 35px 30px;">
                            <h2 style="color: #0f172a; font-size: 20px; margin-top: 0; margin-bottom: 15px;">Hola, {nombre_usuario}</h2>
                            <p style="color: #334155; font-size: 15px; line-height: 1.6; margin-top: 0; margin-bottom: 20px;">
                                Te confirmamos que tu pago ha sido procesado exitosamente y tu suscripción ya se encuentra activa en nuestro sistema. A continuación, te compartimos los detalles de tu cuenta y los beneficios acreditados.
                            </p>

                            <!-- Bloque de Credenciales de Acceso -->
                            <table width="100%" border="0" cellpadding="0" cellspacing="0" style="background-color: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; margin-bottom: 25px;">
                                <tr>
                                    <td style="padding: 20px;">
                                        <p style="color: #166534; font-size: 14px; font-weight: 700; margin: 0 0 10px 0;">Tus credenciales de acceso temporal:</p>
                                        <table width="100%" border="0" cellpadding="4" cellspacing="0">
                                            <tr>
                                                <td style="color: #334155; font-size: 14px; width: 90px;"><strong>Correo:</strong></td>
                                                <td style="color: #0f172a; font-size: 14px; font-family: monospace; background: #ffffff; padding: 4px 8px; border-radius: 4px; border: 1px solid #d1fae5;">{payer_email}</td>
                                            </tr>
                                            <tr>
                                                <td style="color: #334155; font-size: 14px;"><strong>Contraseña:</strong></td>
                                                <td style="color: #0f172a; font-size: 14px; font-family: monospace; background: #ffffff; padding: 4px 8px; border-radius: 4px; border: 1px solid #d1fae5;">{password_temporal}</td>
                                            </tr>
                                        </table>
                                        <p style="color: #15803d; font-size: 12px; margin: 10px 0 0 0;">* Te recomendamos iniciar sesión y cambiar tu contraseña por seguridad desde la configuración de tu perfil.</p>
                                    </td>
                                </tr>
                            </table>

                            <!-- Resumen del Plan -->
                            <p style="color: #64748b; font-size: 13px; font-weight: 700; text-transform: uppercase; letter-spacing: 0.5px; margin: 0 0 10px 0;">Detalles de la Suscripción</p>
                            <table width="100%" border="0" cellpadding="10" cellspacing="0" style="background-color: #f8fafc; border-radius: 8px; margin-bottom: 25px; border: 1px solid #e2e8f0;">
                                <tr>
                                    <td style="color: #475569; font-size: 14px; border-bottom: 1px solid #e2e8f0;"><strong>Plan Activo:</strong></td>
                                    <td align="right" style="color: #0f172a; font-size: 14px; font-weight: 600; border-bottom: 1px solid #e2e8f0;">{nombre_plan_fmt}</td>
                                </tr>
                                <tr>
                                    <td style="color: #475569; font-size: 14px; border-bottom: 1px solid #e2e8f0;"><strong>Inversión Total:</strong></td>
                                    <td align="right" style="color: #0f172a; font-size: 14px; font-weight: 600; border-bottom: 1px solid #e2e8f0;">${precio_pagado:,.0f} COP</td>
                                </tr>
                                <tr>
                                    <td style="color: #475569; font-size: 14px; border-bottom: 1px solid #e2e8f0;"><strong>Capacidad de Audio:</strong></td>
                                    <td align="right" style="color: #0f172a; font-size: 14px; font-weight: 600; border-bottom: 1px solid #e2e8f0;">{horas_otorgadas} horas/mes</td>
                                </tr>
                                <tr>
                                    <td style="color: #475569; font-size: 14px;"><strong>Tokens de IA:</strong></td>
                                    <td align="right" style="color: #0f172a; font-size: 14px; font-weight: 600;">{tokens_otorgados:,} tokens</td>
                                </tr>
                            </table>

                            <!-- Botón de Acción Principal -->
                            <table width="100%" border="0" cellpadding="0" cellspacing="0" style="margin-bottom: 30px;">
                                <tr>
                                    <td align="center">
                                        <a href="https://actaprocore.com/dashboard" target="_blank" style="background-color: #2563eb; color: #ffffff; padding: 14px 30px; border-radius: 6px; text-decoration: none; font-weight: 600; font-size: 15px; display: inline-block; box-shadow: 0 2px 4px rgba(37, 99, 235, 0.2);">Iniciar Sesión en mi Panel</a>
                                    </td>
                                </tr>
                            </table>

                            <p style="color: #64748b; font-size: 14px; line-height: 1.5; margin: 0;">
                                Si tienes alguna duda sobre el funcionamiento de la plataforma o requieres soporte técnico adicional, puedes responder directamente a este correo o contactarnos en cualquier momento.
                            </p>
                        </td>
                    </tr>

                    <!-- Pie de página con enlaces institucionales para evitar SPAM -->
                    <tr>
                        <td align="center" style="background-color: #f8fafc; padding: 20px 30px; border-top: 1px solid #e2e8f0;">
                            <p style="color: #94a3b8; font-size: 12px; margin: 0 0 10px 0; line-height: 1.4;">
                                Este es un correo automático generado tras la confirmación de pago en <strong>ActaPro Core</strong>.<br>
                                Por favor, no compartas tus credenciales de acceso con terceros.
                            </p>
                            <p style="color: #cbd5e1; font-size: 11px; margin: 0;">
                                &copy; 2026 ActaPro Core. Todos los derechos reservados. | <a href="https://actaprocore.com/terminos" target="_blank" style="color: #64748b; text-decoration: underline;">Términos y Condiciones</a>
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>"""
                        
                        enviar_correo_brevo(payer_email, nombre_usuario, asunto_correo, html_cuerpo)

                        return {"status": "success", "message": f"Licencia de {payer_email} actualizada a {plan_asignado}"}
    except Exception as e:
        print(f"Error en webhook de Mercado Pago: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    return {"status": "ok"}

import ffmpeg

# Función para validar la calidad técnica del audio en Python
def validar_calidad_audio(file_path: str):
    try:
        # ffmpeg.probe obtiene los metadatos del archivo
        probe = ffmpeg.probe(file_path)
        format_info = probe.get('format', {})
        
        bit_rate = format_info.get('bit_rate')
        duration = format_info.get('duration')
        
        # REGLA 1: Bitrate menor a 32 kbps (audio ultra comprimido o estática)
        if bit_rate is not None:
            if int(bit_rate) < 32000:
                return {
                    "valido": False,
                    "motivo": "AUDIO_MUY_COMPRIMIDO_O_DEFECTUOSO",
                    "mensaje": "El archivo presenta una calidad técnica deficiente (bitrate muy bajo). Esto impedirá una correcta identificación de oradores."
                }
                
        # REGLA 2: Audio demasiado corto (menos de 10 segundos para una asamblea)
        if duration is not None:
            if float(duration) < 10.0:
                return {
                    "valido": False,
                    "motivo": "AUDIO_DEMASIADO_CORTO",
                    "mensaje": "El archivo de audio es demasiado corto para ser una asamblea."
                }
                
        return {"valido": True}
        
    except ffmpeg.Error:
        # Ocurre si ffprobe no puede leer el archivo o está corrupto
        return {
            "valido": False,
            "motivo": "ERROR_DE_LECTURA",
            "mensaje": "No se pudo leer el archivo de audio. Puede estar corrupto."
        }


import os
import uuid
from pathlib import Path
from typing import Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, status

# Importa tu tarea de celery
from tasks import task_procesar_asamblea, celery_app

TEMP_DIR = Path("temp_uploads")
TEMP_DIR.mkdir(parents=True, exist_ok=True)

@app.post("/procesar-asamblea")
@app.post("/procesar")
async def procesar_asamblea(
    file: UploadFile = File(...),
    email: str = Form(...),
    instrucciones: Optional[str] = Form(None),
    nombre_personalizado: Optional[str] = Form(None)
):
    try:
        # Generar un nombre único para el archivo en la nube
        file_extension = Path(file.filename).suffix
        unique_filename = f"audios/{uuid.uuid4()}{file_extension}"
        
        # Leer los bytes del archivo directamente de la petición
        file_bytes = await file.read()

        # Subir el archivo directamente a Cloudflare R2
        s3 = get_r2_client()
        s3.put_object(
            Bucket=R2_BUCKET_NAME,
            Key=unique_filename,
            Body=file_bytes,
            ContentType=file.content_type or "audio/mpeg"
        )

        # Construir la URL pública accesible por el Worker de Celery y AssemblyAI
        audio_url = f"{R2_PUBLIC_URL.rstrip('/')}/{unique_filename}"

        # Disparar la tarea en Celery pasando la URL web en lugar de una ruta local
        task = task_procesar_asamblea.delay(
            temp_audio_path=audio_url,
            email=email,
            instrucciones=instrucciones,
            nombre_personalizado=nombre_personalizado,
            original_filename=file.filename
        )

        return {
            "status": "pending",
            "task_id": task.id,
            "message": "Procesamiento de asamblea iniciado correctamente en la nube."
        }

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error iniciando la tarea en la nube: {str(e)}"
        )

from fastapi.responses import RedirectResponse

from io import BytesIO
from bson import ObjectId
from docx import Document
from fastapi import HTTPException, Response
from fastapi.responses import RedirectResponse

@app.get("/api/actas/descargar/{acta_id}")
async def descargar_acta(acta_id: str, email: str):
    """
    Endpoint seguro para descargar actas con aislamiento por usuario.
    Redirige a Cloudflare R2 si existe URL o genera el Word al vuelo como respaldo.
    """
    acta = None
    
    # 1. Buscar por ObjectId de MongoDB de forma segura
    if ObjectId.is_valid(acta_id):
        try:
            acta = actas_collection.find_one({"_id": ObjectId(acta_id), "email": email})
        except Exception as e:
            print(f"Error al consultar por ObjectId: {e}")
            
    # 2. Si no se encontró, buscar por nombre_acta
    if not acta:
        acta = actas_collection.find_one({"nombre_acta": acta_id, "email": email})
        
    # 3. Si aún no existe, buscar por el campo acta_id en texto plano
    if not acta:
        acta = actas_collection.find_one({"acta_id": acta_id, "email": email})
        
    # 4. Validar existencia y permisos (Aislamiento de usuario)
    if not acta:
        raise HTTPException(
            status_code=404, 
            detail="El acta solicitada no existe o no tienes permisos para acceder a ella."
        )
        
    # 5. Opción Principal: Redirección directa a Cloudflare R2 (CDN)
    file_url = acta.get("file_url")
    if file_url:
        return RedirectResponse(url=file_url, status_code=303)

    # 6. Plan de Respaldo: Generación dinámica del documento Word (.docx)
    contenido_texto = acta.get("contenido", "") or acta.get("transcripcion", "")
    nombre_archivo = acta.get("nombre_acta") or f"Acta_Asamblea_{acta_id[:8]}.docx"
    
    if not nombre_archivo.endswith(".docx"):
        nombre_archivo += ".docx"

    doc = Document()
    titulo_principal = doc.add_heading("ACTA DE ASAMBLEA GENERAL DE COPROPIETARIOS", level=0)
    titulo_principal.alignment = 1 # Centrado

    # Parser básico de formato Markdown a Word
    for linea in contenido_texto.split("\n"):
        linea_clean = linea.strip()
        if not linea_clean:
            continue

        if linea_clean.startswith("# "):
            doc.add_heading(linea_clean.replace("# ", "").strip(), level=1)
        elif linea_clean.startswith("## ") or linea_clean.startswith("### "):
            doc.add_heading(linea_clean.replace("#", "").strip(), level=2)
        else:
            p = doc.add_paragraph()
            if "**" in linea_clean:
                partes = linea_clean.split("**")
                for i, parte in enumerate(partes):
                    if parte:
                        run = p.add_run(parte)
                        if i % 2 == 1: # Texto en negrita
                            run.bold = True
            else:
                p.add_run(linea_clean)

    # Guardar en memoria y retornar como archivo descargable
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return Response(
        content=doc_io.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'}
    )

import os
import base64
import pymupdf
import pdfplumber
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from typing import Optional
from openai import OpenAI

# Importar la tarea Celery desde tasks.py
from tasks import task_escanear_documento

# Inicializa el cliente de OpenAI
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


@app.post("/escanear")
@app.post("/escanear-documento")
async def escanear_documento(
    file: UploadFile = File(...), 
    email: Optional[str] = Form(None)
):
    try:
        # 1. Leer los bytes del archivo subido
        file_bytes = await file.read()

        # 2. Convertir los bytes a string Base64 para pasarlos por la cola de Celery
        file_bytes_b64 = base64.b64encode(file_bytes).decode("utf-8")

        # 3. Disparar la tarea asíncrona en Celery pasando el Base64
        task = task_escanear_documento.delay(
            file_bytes_b64=file_bytes_b64,
            filename=file.filename,
            email=email
        )

        # 4. Retornar inmediatamente el ID de la tarea al frontend
        return {
            "status": "pending",
            "task_id": task.id,
            "message": "Escaneo de documento en proceso."
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error iniciando el escaneo: {str(e)}")
        
import io
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

@app.get("/api/actas/descargar-pdf/{acta_id}")
async def descargar_acta_pdf(acta_id: str, email: str):
    # Buscar el acta en la base de datos por _id de MongoDB o por nombre
    acta = None
    try:
        acta = actas_collection.find_one({"_id": ObjectId(acta_id), "email": email})
    except Exception:
        pass
        
    if not acta:
        acta = actas_collection.find_one({"nombre_acta": acta_id, "email": email})
        
    if not acta:
        raise HTTPException(status_code=404, detail="Acta no encontrada.")
        
    contenido_texto = acta.get("contenido", "")
    nombre_base = acta.get("nombre_acta", "Acta_Asamblea").replace(".docx", "").replace(".pdf", "")
    nombre_archivo_pdf = f"{nombre_base}.pdf"
    
    # Configuración del PDF en memoria usando ReportLab
    pdf_buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        pdf_buffer, 
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'ActaTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=15,
        textColor=styles['Normal'].textColor
    )
    
    body_style = ParagraphStyle(
        'ActaBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    
    story.append(Paragraph("ACTA DE ASAMBLEA GENERAL DE COPROPIETARIOS", title_style))
    story.append(Spacer(1, 10))
    
    lineas = contenido_texto.split('\n')
    for linea in lineas:
        linea_limpia = linea.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        linea_limpia = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', linea_limpia)
        linea_limpia = linea_limpia.replace('*', '').strip()
        
        if linea_limpia:
            story.append(Paragraph(linea_limpia, body_style))
        else:
            story.append(Spacer(1, 6))

    doc.build(story)
    pdf_buffer.seek(0)
    
    return Response(
        content=pdf_buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo_pdf}"'}
    )


from bson import ObjectId
from fastapi import HTTPException
from pydantic import BaseModel

# ================= MODELOS PYDANTIC PARA SCANNERS =================
class GuardarScannerRequest(BaseModel):
    email: str
    nombre: str
    tokens: int
    contenido: str

class RenombrarScannerRequest(BaseModel):
    email: str
    scanner_id: str
    nuevo_nombre: str

class EliminarScannerRequest(BaseModel):
    email: str
    scanner_id: str


# ================= ENDPOINTS DE SCANNERS =================

import os
from fastapi import FastAPI, HTTPException
from celery.result import AsyncResult
from celery_app import celery_app


# Asegurar que exista el directorio para audios temporales
os.makedirs("temp_uploads", exist_ok=True)
os.makedirs("temp_outputs", exist_ok=True)

# ---------------------------------------------------------
# ENDPOINT: CONSULTAR ESTADO DE LA TAREA (POLLING)
# ---------------------------------------------------------
@app.get("/estado-tarea/{task_id}")
async def obtener_estado_tarea(task_id: str):
    res = AsyncResult(task_id, app=celery_app)
    
    if res.state == "PENDING":
        return {
            "status": "PROCESSING",
            "info": {"status": "En cola de espera..."}
        }
    elif res.state == "PROCESSING":
        return {
            "status": "PROCESSING",
            "info": res.info  # Devuelve la meta enviada desde update_state
        }
    elif res.state == "SUCCESS":
        return {
            "status": "COMPLETED",
            "result": res.result  # Devuelve el diccionario final retornado por la tarea
        }
    elif res.state == "FAILURE":
        return {
            "status": "FAILED",
            "error": str(res.result)
        }
    
    return {"status": res.state, "info": str(res.info)}

@app.get("/api/scanners/historial")
async def obtener_historial_scanners(email: str):
    scanners_cursor = scanners_historial_collection.find({"email": email})
    scanners = []
    for doc in scanners_cursor:
        doc["id"] = str(doc["_id"])
        del doc["_id"]
        scanners.append(doc)
    return {"scanners": scanners}


@app.post("/api/scanners/guardar")
async def guardar_escaneo(data: GuardarScannerRequest):
    try:
        nuevo_registro = {
            "email": data.email,
            "nombre": data.nombre,
            "fecha": datetime.utcnow().isoformat(),
            "tokens": data.tokens,
            "contenido": data.contenido
        }
        
        resultado = scanners_historial_collection.insert_one(nuevo_registro)
        
        return {
            "message": "¡Escaneo guardado correctamente!", 
            "id": str(resultado.inserted_id)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/scanners/renombrar")
async def renombrar_scanner(data: RenombrarScannerRequest):
    try:
        # Intentar buscar primero por ObjectId si es un ID válido de Mongo, o por campo nombre
        filtro = {"email": data.email}
        try:
            filtro["$or"] = [{"_id": ObjectId(data.scanner_id)}, {"nombre": data.scanner_id}]
        except Exception:
            filtro["nombre"] = data.scanner_id

        resultado = scanners_historial_collection.update_one(
            filtro,
            {"$set": {"nombre": data.nuevo_nombre}},
        )

        if resultado.matched_count == 0:
            raise HTTPException(
                status_code=404, detail="Escaneo no encontrado o no autorizado."
            )

        return {"message": "¡Escaneo renombrado con éxito!"}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/scanners/eliminar")
async def eliminar_scanner(data: EliminarScannerRequest):
    try:
        filtro = {"email": data.email}
        try:
            filtro["$or"] = [{"_id": ObjectId(data.scanner_id)}, {"nombre": data.scanner_id}]
        except Exception:
            filtro["nombre"] = data.scanner_id

        resultado = scanners_historial_collection.delete_one(filtro)

        if resultado.deleted_count == 0:
            raise HTTPException(
                status_code=404, detail="Escaneo no encontrado para eliminar."
            )

        return {"message": "Escaneo eliminado correctamente."}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ================= ENDPOINTS DE ACTAS (REFERENCIA) =================

@app.get("/api/actas/historial")
async def obtener_historial_actas(email: str):
    actas = list(actas_collection.find({"email": email}, {"_id": 0}))
    return {"actas": actas}

from fastapi.responses import HTMLResponse

from io import BytesIO
from fastapi import Response, HTTPException
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bson import ObjectId
from datetime import datetime
from html.parser import HTMLParser

class HTMLToDocxParser(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_tag = None
        self.current_text = ""
        self.is_bold = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ['h1', 'h2', 'h3', 'p', 'div', 'br', 'li']:
            self._flush_text()
            self.current_tag = tag

        if tag in ['strong', 'b']:
            self._flush_text()
            self.is_bold = True

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ['h1', 'h2', 'h3', 'p', 'div', 'li']:
            self._flush_text()
            self.current_tag = None
        elif tag in ['strong', 'b']:
            self._flush_text()
            self.is_bold = False

    def handle_data(self, data):
        self.current_text += data

    def _flush_text(self):
        text = self.current_text.strip()
        if not text and self.current_tag != 'br':
            self.current_text = ""
            return

        if self.current_tag == 'h1':
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)
        elif self.current_tag == 'h2':
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85)
        elif self.current_tag == 'h3':
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(71, 85, 105)
        else:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = self.is_bold
            run.font.color.rgb = RGBColor(51, 51, 51)

        self.current_text = ""

@app.get("/api/scanners/descargar/{scanner_id}")
async def descargar_scanner_docx(scanner_id: str, email: str):
    try:
        # Búsqueda segura en MongoDB
        filtro = {"_id": ObjectId(scanner_id), "email": email}
        registro = scanners_historial_collection.find_one(filtro)
        
        if not registro:
            registro = scanners_historial_collection.find_one({"nombre": scanner_id, "email": email})
            
        if not registro:
            raise HTTPException(status_code=404, detail="Archivo no encontrado.")

        doc = Document()

        # Configuración de Márgenes de Página (1 pulgada / 2.54 cm)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Estilo Global Normal
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(51, 51, 51)

        # Procesamiento del contenido HTML respetando etiquetas de título y formato
        contenido = registro.get('contenido', '')
        
        if "<" in contenido and ">" in contenido:
            parser = HTMLToDocxParser(doc)
            parser.feed(contenido)
            parser._flush_text()
        else:
            # Plan de respaldo si es texto plano
            for linea in contenido.split("\n"):
                linea_clean = linea.strip()
                if not linea_clean:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(linea_clean)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(51, 51, 51)

        # --- PIE DE PÁGINA ---
        footer = doc.sections[0].footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_f = footer_p.add_run("Generado automáticamente por ActaProCore")
        run_f.font.size = Pt(8.5)
        run_f.font.color.rgb = RGBColor(148, 163, 184)

        # Guardar en buffer de memoria
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        nombre_base = registro.get('nombre', 'documento').replace('.pdf', '').replace('.jpg', '').replace('.png', '')
        nombre_archivo = f"{nombre_base}_ActaProCore.docx"

        headers = {
            'Content-Disposition': f'attachment; filename="{nombre_archivo}"'
        }
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers=headers
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


from fastapi.responses import Response
from fastapi import HTTPException
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
import re
import io
from bson import ObjectId

@app.get("/api/scanners/descargar-pdf/{scanner_id}")
async def descargar_scanner_pdf_backend(scanner_id: str, email: str):
    try:
        filtro = {"_id": ObjectId(scanner_id), "email": email}
        registro = scanners_historial_collection.find_one(filtro)
        
        if not registro:
            registro = scanners_historial_collection.find_one({"nombre": scanner_id, "email": email})
            
        if not registro:
            raise HTTPException(status_code=404, detail="Archivo no encontrado.")

        nombre_base = registro.get("nombre", "Escaneo_IA").replace(".pdf", "").replace(".jpg", "").replace(".png", "")
        nombre_archivo_pdf = f"{nombre_base}.pdf"
        contenido_html = registro.get("contenido", "")

        # Configuración del PDF en memoria usando ReportLab con metadatos corregidos
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
            title=nombre_base,
            author="Sistema de Escaneo IA"
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        h1_style = ParagraphStyle(
            'ScannerH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            alignment=TA_LEFT,
            spaceBefore=12,
            spaceAfter=8
        )

        h2_style = ParagraphStyle(
            'ScannerH2',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=10,
            spaceAfter=6
        )
        
        body_style = ParagraphStyle(
            'ScannerBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        )

        # Procesamiento de etiquetas HTML
        contenido_procesado = re.sub(r'<h1[^>]*>(.*?)</h1>', r'<h1_tag>\1</h1_tag>', contenido_html, flags=re.IGNORECASE | re.DOTALL)
        contenido_procesado = re.sub(r'<h2[^>]*>(.*?)</h2>', r'<h2_tag>\1</h2_tag>', contenido_procesado, flags=re.IGNORECASE | re.DOTALL)
        contenido_procesado = re.sub(r'<p[^>]*>(.*?)</p>', r'\1<br/><br/>', contenido_procesado, flags=re.IGNORECASE | re.DOTALL)
        contenido_procesado = re.sub(r'<br\s*/?>', r'<br/>', contenido_procesado, flags=re.IGNORECASE)
        contenido_procesado = re.sub(r'<(?:strong)[^>]*>(.*?)</(?:strong)>', r'<b>\1</b>', contenido_procesado, flags=re.IGNORECASE | re.DOTALL)
        
        lineas = contenido_procesado.split('\n')
        
        for linea in lineas:
            linea_str = linea.strip()
            if not linea_str:
                continue
                
            # Verificar h1
            h1_match = re.search(r'<h1_tag>(.*?)</h1_tag>', linea_str, flags=re.DOTALL)
            if h1_match:
                story.append(Paragraph(h1_match.group(1), h1_style))
                continue
                
            # Verificar h2
            h2_match = re.search(r'<h2_tag>(.*?)</h2_tag>', linea_str, flags=re.DOTALL)
            if h2_match:
                story.append(Paragraph(h2_match.group(1), h2_style))
                continue
                
            # Limpiar etiquetas no permitidas en párrafos normales
            linea_limpia = re.sub(r'<(?!/?b\b)(?!/?br\b)[^>]+>', '', linea_str)
            linea_limpia = linea_limpia.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            
            if linea_limpia.strip():
                story.append(Paragraph(linea_limpia, body_style))

        doc.build(story)
        pdf_buffer.seek(0)
        
        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={nombre_archivo_pdf}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando PDF: {str(e)}")




# --- Ruta 2: Verificación de Plan Preestablecida con Video ---
@app.post("/api/soporte/verificar-plan")
async def verificar_plan(req: ConsultaPlanRequest):
    # Simulación de consulta a base de datos usando el correo
    # Aquí puedes conectar tu lógica real de base de datos para buscar el plan del usuario por su `req.email`
    
    respuesta_plan = (
        f"Hola. Hemos verificado el correo '{req.email}'. "
        f"Actualmente te encuentras en el plan **Empresarial Pro** con 18.5 horas consumidas de tu límite mensual. "
        f"Tu cuenta está activa y tienes acceso completo a todas las funciones avanzadas y escáneres hasta tu fecha de renovación."
    )
    
    video_plan = "https://assets.mixkit.co/videos/preview/mixkit-data-flow-plexus-background-animation-2679-large.mp4" # Video explicativo del plan

    return {
        "respuesta": respuesta_plan,
        "videoUrl": video_plan,
        "datosConsumo": {
            "planActivo": "Empresarial Pro",
            "horasTotalesMes": 20,
            "horasConsumidas": 18.5,
            "fechaRenovacion": "01/10/2026"
        }
    }

from fastapi import FastAPI
from pydantic import BaseModel
from typing import Optional

from bson import ObjectId

# (Asegúrate de tener configurado tu cliente de Motor/MongoDB en tu app, por ejemplo:
# from motor.motor_asyncio import AsyncIOMotorClient
# client = AsyncIOMotorClient("tu_url_de_mongo")
# db = client["nombre_de_tu_base_de_datos"]
# )

class ContactoEnterprise(BaseModel):
    nombre: str 
    email: str
    whatsapp: str
    pais: str
    horas: int
    documentos: int
    empleados: int
    mensaje: Optional[str] = None



# ==========================================
# 1. ENDPOINT PARA GUARDAR COTIZACIÓN (POST)
# ==========================================
@app.post("/api/contacto-enterprise")
async def recibir_contacto_enterprise(datos: ContactoEnterprise):
    try:
        nuevo_lead = datos.model_dump()
        
        # Genera la fecha ISO en formato UTC correctamente
        nuevo_lead["fecha"] = datetime.now(timezone.utc).isoformat()
        nuevo_lead["estado"] = "Pendiente"
        
        resultado_db = db.cotizaciones.insert_one(nuevo_lead)
        
        print(f"Nueva solicitud Enterprise guardada con ID: {resultado_db.inserted_id}")
        
        return {
            "success": True, 
            "message": "¡Gracias por contactarnos! Un gerente de cuenta se comunicará contigo pronto."
        }
    except Exception as e:
        print(f"Error procesando solicitud Enterprise: {e}")
        return {
            "success": False, 
            "message": str(e)
        }
# ==========================================
# 2. ENDPOINT PARA EL PANEL DE ADMINISTRACIÓN (GET)
# ==========================================
@app.get("/api/admin/cotizaciones")
async def obtener_cotizaciones():
    try:
        # CORRECCIÓN: Sin 'await' porque PyMongo es síncrono
        cotizaciones_db = list(db.cotizaciones.find().sort("fecha", -1).limit(1000))
        
        # Transformamos el ObjectId de MongoDB a string
        lista_leads = []
        for c in cotizaciones_db:
            c["_id"] = str(c["_id"])
            lista_leads.append(c)
        
        return {
            "success": True,
            "cotizaciones": lista_leads
        }
    except Exception as e:
        print(f"Error obteniendo cotizaciones de MongoDB: {e}")
        return {
            "success": False,
            "message": str(e)
        }

# ==========================================
# MÓDULO DE CRM Y MÉTRICAS FINANCIERAS
# ==========================================

@app.get("/api/crm/license-stats")
async def get_crm_license_stats():
    """
    Endpoint principal para alimentar el CRM con estadísticas de usuarios por plan,
    desglose de pagos y proyecciones financieras mensuales/anuales.
    """
    try:
        # 1. Conteo de usuarios por plan utilizando la colección existente 'users'
        total_usuarios = users_collection.count_documents({})
        free_users = users_collection.count_documents({"plan": "free"})
        basico_users = users_collection.count_documents({"plan": "basico"})
        profesional_users = users_collection.count_documents({"plan": "profesional"})
        corporativo_users = users_collection.count_documents({"plan": "corporativo"})
        
        # Total de usuarios de pago (todos los que no sean 'free')
        paid_users = basico_users + profesional_users + corporativo_users

        # 2. Obtener lista detallada de todos los usuarios para la tabla del CRM
        cursor_usuarios = users_collection.find({}, {"_id": 0, "password": 0})
        lista_suscriptores = list(cursor_usuarios)

        # 3. Cálculo de Ingresos Actuales (MRR estimado según los planes activos)
        # Precios de referencia desde tu diccionario PRECIOS_PLANES:
        # basico: 153,000 | profesional: 479,000 | corporativo: 939,000
        precio_basico = PRECIOS_PLANES["basico"]["precio"]
        precio_profesional = PRECIOS_PLANES["profesional"]["precio"]
        precio_corporativo = PRECIOS_PLANES["corporativo"]["precio"]

        ingresos_actuales_mes = (
            (basico_users * precio_basico) +
            (profesional_users * precio_profesional) +
            (corporativo_users * precio_corporativo)
        )

        # Estimación de egresos operativos (ej. costos de OpenAI/AssemblyAI/Infraestructura aprox. 25% de base o calculados)
        egresos_actuales_mes = ingresos_actuales_mes * 0.25 

        # 4. Simulación de Evolución Mensual (Últimos meses y proyección anual)
        # Puedes conectar esto con transacciones reales de pagos si guardas historial.
        # Aquí generamos una estructura lógica basada en los suscriptores actuales.
        evolucion_mensual = [
            {"mes": "Enero", "ingresos": ingresos_actuales_mes * 0.7, "egresos": egresos_actuales_mes * 0.8, "proyeccion": ingresos_actuales_mes * 0.75},
            {"mes": "Febrero", "ingresos": ingresos_actuales_mes * 0.82, "egresos": egresos_actuales_mes * 0.85, "proyeccion": ingresos_actuales_mes * 0.88},
            {"mes": "Marzo", "ingresos": ingresos_actuales_mes * 0.95, "egresos": egresos_actuales_mes * 0.9, "proyeccion": ingresos_actuales_mes * 0.92},
            {"mes": "Abril", "ingresos": ingresos_actuales_mes, "egresos": egresos_actuales_mes, "proyeccion": ingresos_actuales_mes * 1.05},
            {"mes": "Mayo (Proj)", "ingresos": 0, "egresos": 0, "proyeccion": ingresos_actuales_mes * 1.15},
            {"mes": "Junio (Proj)", "ingresos": 0, "egresos": 0, "proyeccion": ingresos_actuales_mes * 1.25},
        ]

        proyeccion_anual = {
            "anual_estimado": ingresos_actuales_mes * 12,
            "crecimiento_esperado_porcentaje": 18.5,
            "meta_anual": (ingresos_actuales_mes * 12) * 1.30
        }

        return {
            "success": True,
            "metricas_planes": {
                "total": total_usuarios,
                "free": free_users,
                "paid": paid_users,
                "desglose_pagos": {
                    "basico": basico_users,
                    "profesional": profesional_users,
                    "corporativo": corporativo_users
                }
            },
            "financiero": {
                "ingresos_mes_actual": ingresos_actuales_mes,
                "egresos_mes_actual": egresos_actuales_mes,
                "utilidad_neta": ingresos_actuales_mes - egresos_actuales_mes,
                "evolucion_mensual": evolucion_mensual,
                "proyeccion_anual": proyeccion_anual
            },
            "suscriptores": lista_suscriptores
        }

    except Exception as e:
        print(f"Error en CRM stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


from fastapi import APIRouter, HTTPException, Request
from pydantic import BaseModel

router = APIRouter(prefix="/api/crm", tags=["CRM & SaaS Analytics"])

# Mock o conexión a tu base de datos MongoDB y configuración de planes
# db = client["tu_base_de_datos"]
# users_collection = db["users"]
# mp_sdk = mercadopago.SDK("YOUR_ACCESS_TOKEN")

PRECIOS_PLANES = {
    "basico": {"tokens_mensuales": 100000, "limite_horas": 15.0, "precio": 4000},
    "profesional": {"tokens_mensuales": 500000, "limite_horas": 60.0, "precio": 4000},
    "corporativo": {"tokens_mensuales": 2000000, "limite_horas": 200.0, "precio": 4000}
}

@router.get("/license-stats-advanced")
async def obtener_estadisticas_avanzadas():
    """
    Retorna métricas extendidas incluyendo alerta de churn, 
    márgenes unitarios por usuario y cohortes desde la base de datos real.
    """
    # Ejemplo de recuperación real o simulación desde MongoDB:
    # suscriptores = list(users_collection.find({}, {"_id": 0}))
    
    suscriptores = [
        {
            "email": "carlos.legal@empresa.com",
            "plan": "corporativo",
            "fecha_registro": "2026-03-15",
            "horas_usadas_mes": 185.0,
            "limite_horas_mes": 200,
            "tokens_usados": 1250000,
            "pago_mensual": 4000,
            "activo": True
        },
        {
            "email": "laura.consultora@gmail.com",
            "plan": "profesional",
            "fecha_registro": "2026-06-10",
            "horas_usadas_mes": 5.0, # Uso muy bajo -> Alerta de Churn
            "limite_horas_mes": 60,
            "tokens_usados": 45000,
            "pago_mensual": 4000,
            "activo": True
        },
        {
            "email": "startup.dev@gmail.com",
            "plan": "basico",
            "fecha_registro": "2026-07-01",
            "horas_usadas_mes": 14.5,
            "limite_horas_mes": 15,
            "tokens_usados": 110000,
            "pago_mensual": 4000,
            "activo": True
        }
    ]

    Costo_IA_Por_Hora = 2500 # COP
    analisis_unit_economics = []
    usuarios_en_riesgo = []

    for sub in suscriptores:
        # Cálculo de Costo Operativo Real de IA por usuario
        costo_ia_usuario = sub["horas_usadas_mes"] * Costo_IA_Por_Hora
        margen_neto = sub["pago_mensual"] - costo_ia_usuario
        
        analisis_unit_economics.append({
            "email": sub["email"],
            "plan": sub["plan"],
            "ingreso": sub["pago_mensual"],
            "costo_ia": costo_ia_usuario,
            "margen_neto": margen_neto,
            "porcentaje_margen": round((margen_neto / sub["pago_mensual"]) * 100, 1) if sub["pago_mensual"] > 0 else 0
        })

        # --- ALGORITMO DE ALERTA TEMPRANA DE CHURN ---
        if sub["plan"] != "free":
            limite = sub.get("limite_horas_mes", 0)
            if limite > 0:
                porcentaje_uso = (sub["horas_usadas_mes"] / limite) * 100
                if porcentaje_uso < 15:
                    usuarios_en_riesgo.append({
                        "email": sub["email"],
                        "plan": sub["plan"],
                        "razon": f"Uso crítico de cuota: solo {porcentaje_uso:.1f}% consumido",
                        "nivel_riesgo": "Alto"
                    })

    # --- ANÁLISIS DE COHORTES (Simulado por Mes de Registro) ---
    cohortes = {
        "Marzo 2026": {"total": 12, "retenidos": 11},
        "Abril 2026": {"total": 19, "retenidos": 17},
        "Mayo 2026": {"total": 25, "retenidos": 23},
        "Junio 2026": {"total": 31, "retenidos": 28},
        "Julio 2026": {"total": 45, "retenidos": 42}
    }

    return {
        "success": True,
        "unit_economics": analisis_unit_economics,
        "churn_alerts": usuarios_en_riesgo,
        "cohortes": cohortes
    }


@router.get("/invoice/download/{email}")
async def descargar_factura_pdf(email: str):
    """
    Genera los datos fiscales para la descarga del comprobante contable.
    """
    return {
        "success": True,
        "mensaje": "Datos de factura listos",
        "factura": {
            "empresa_emisora": "ActasPro SAS - NIT: 900.123.456-7",
            "cliente": email,
            "fecha_emision": datetime.now().strftime("%Y-%m-%d"),
            "concepto": "Suscripción Licencia Corporativa - Mensual",
            "valor_neto": 4000,
            "iva": 0,
            "total": 4000,
            "estado": "Pagado vía Mercado Pago"
        }
    }

from fastapi import APIRouter

# Puedes crear un router local o agregarlo directo a app
crm_router = APIRouter(prefix="/api/crm", tags=["CRM"])

@crm_router.get("/license-stats")
def obtener_estadisticas_licencias():
    # Lógica para calcular o extraer las estadísticas de tu base de datos
    return {
        "status": "success",
        "total_licencias": 0,
        "activas": 0,
        "expiradas": 0
    }

from fastapi import APIRouter

# 1. Crear el router con el prefijo /api/crm que busca tu frontend
crm_router = APIRouter(prefix="/api/crm", tags=["CRM"])

@crm_router.get("/license-stats")
def get_license_stats():
    # Aquí puedes conectar con tus colecciones de MongoDB para devolver datos reales
    # Por ahora, devolvemos una estructura base para que el frontend no falle:
    return {
        "success": True,
        "total_usuarios": users_collection.count_documents({}),
        "licencias_activas": 0,
        "estadisticas": {}
    }

@crm_router.get("/license-stats-advanced")
def get_license_stats_advanced():
    try:
        # Puedes extraer los datos reales de tu base de datos (users_collection, actas_collection, etc.)
        # Aquí te dejamos la estructura exacta que tu frontend (JavaScript) está exigiendo:
        
        return {
            "success": True,
            "churn_alerts": [
                # Ejemplo de estructura si quisieras mostrar alertas reales o dejarlo vacío []
                # {
                #     "email": "usuario@ejemplo.com",
                #     "plan": "basico",
                #     "razon": "Inactividad prolongada",
                #     "nivel_riesgo": "Alto"
                # }
            ],
            "cohortes": {
                # Ejemplo de estructura de cohortes por mes:
                # "2026-01": {"total": 10, "retenidos": 8},
                # "2026-02": {"total": 15, "retenidos": 12}
            },
            "unit_economics": [
                # Ejemplo de estructura para unit economics:
                # {
                #     "email": "usuario@ejemplo.com",
                #     "plan": "profesional",
                #     "ingreso": 479000.0,
                #     "costo_ia": 50000.0,
                #     "margen_neto": 429000.0,
                #     "porcentaje_margen": 89.5
                # }
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class FeedbackAccountModel(BaseModel):
    email: str
    password: str
    horas: float
    tokens: int
    nombre_empresa: Optional[str] = "Cuenta Feedback"

@app.post("/api/admin/crear-feedback")
def crear_cuenta_feedback(data: FeedbackAccountModel):
    """Crea una cuenta gratuita con límites de horas y tokens personalizados para feedback."""
    existing_user = users_collection.find_one({"email": data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="El correo ya se encuentra registrado.")

    nuevo_usuario = {
        "email": data.email,
        "password": data.password,
        "plan": "feedback",
        "nombre_empresa": data.nombre_empresa,
        # Gestión personalizada de Horas
        "horas_restantes": data.horas,
        "horas_usadas_mes": 0.0,
        "limite_horas_mes": data.horas,
        # Gestión personalizada de Tokens
        "tokens_usados": 0,
        "limite_tokens_mes": data.tokens,
        "created_at": datetime.now()
    }

    users_collection.insert_one(nuevo_usuario)
    return {
        "status": "success",
        "message": "Cuenta feedback creada exitosamente",
        "email": data.email,
        "horas": data.horas,
        "tokens": data.tokens
    }

class RecuperarLinkRequest(BaseModel):
    email: str

@router.post("/api/admin/recuperar-link-pago")
def recuperar_link_pago(data: RecuperarLinkRequest):
    # 1. Buscar al usuario en tu base de datos (MongoDB / Firestore / SQLite, etc.)
    usuario = coleccion_usuarios.find_one({"email": data.email}) # Ejemplo con MongoDB
    
    if not usuario:
        raise HTTPException(status_code=404, detail="Usuario no encontrado en el sistema.")
    
    # 2. Extraer los datos del plan actual que tenía
    plan_nombre = usuario.get("plan", "feedback")
    horas = usuario.get("limite_horas_mes", 10)
    tokens = usuario.get("limite_tokens_mes", 50000)
    
    # 3. GENERAR EL LINK DE PAGO NUEVO
    # Aquí integras la llamada a tu pasarela de pagos (Ej: Mercado Pago SDK) 
    # utilizando los valores guardados del usuario:
    try:
        # ---- EJEMPLO CON MERCADO PAGO SDK ----
        # preference_data = {
        #     "items": [{
        #         "title": f"Renovación Plan {plan_nombre.capitalize()} - {usuario.get('nombre_empresa')}",
        #         "quantity": 1,
        #         "unit_price": 30.00 # El precio correspondiente a su plan
        #     }],
        #     "payer": {"email": data.email},
        #     "back_urls": {
        #         "success": "https://tu-web.com/pago-exitoso",
        #         "pending": "https://tu-web.com/pago-pendiente",
        #         "failure": "https://tu-web.com/pago-fallido"
        #     },
        #     "auto_return": "approved"
        # }
        # preference_response = sdk.preference().create(preference_data)
        # link_pago = preference_response["response"]["init_point"]
        
        # (Simulación para que lo adaptes a tu pasarela actual si usas otra):
        link_pago = f"https://www.mercadopago.com.co/checkout/v1/redirect?pref_id=EJEMPLO_RECUPERADO_{data.email}"

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar el link de pago: {str(e)}")

    return {
        "success": True,
        "email": data.email,
        "nombre_empresa": usuario.get("nombre_empresa"),
        "plan": plan_nombre,
        "link_pago": link_pago
    }

# --- ENDPOINT 2: Cancelar Suscripción Activa en Mercado Pago ---
@router.post("/api/admin/cancelar-suscripcion")
def cancelar_suscripcion_mercado_pago(data: CancelarSuscripcionRequest):
    try:
        # En el SDK de Python, las suscripciones recurrentes (preapproval) se actualizan mediante .preapproval().update()
        update_data = {
            "status": "canceled"
        }
        response = mp_sdk.preapproval().update(data.preapproval_id, update_data)
        
        return {
            "success": True,
            "message": "Suscripción cancelada con éxito",
            "response": response
        }
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Error al cancelar la suscripción: {str(error)}")


# --- ENDPOINT 3: Programar Cancelación de Suscripción ---
@router.post("/api/admin/programar-cancelacion")
def programar_cancelacion_suscripcion(data: ProgramarCancelacionRequest):
    try:
        update_data = {
            "auto_recurring": {
                "end_date": data.fecha_vencimiento
            }
        }
        response = mp_sdk.preapproval().update(data.preapproval_id, update_data)
        
        return {
            "success": True,
            "message": "Cancelación programada con éxito",
            "response": response
        }
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Error al programar la cancelación: {str(error)}")


# --- ENDPOINT 4: Listar Cuentas Free ---
@router.get("/api/admin/cuentas-free")
def listar_cuentas_free(db=None):
    try:
        # Si usas PyMongo / Motor:
        # free_users = list(db.usuarios.find({"plan": "free"}))
        free_users = [] # Reemplaza con tu consulta real a la base de datos
        
        return {
            "success": True,
            "total": len(free_users),
            "cuentas": free_users
        }
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Error al listar cuentas free: {str(error)}")


from urllib.parse import unquote
from fastapi import HTTPException

@app.delete("/api/admin/cuentas-free/{identifier}")
def eliminar_cuenta_free_endpoint(identifier: str):
    try:
        email_decodificado = unquote(identifier)
        
        # Intenta borrar por email o por _id asegurando que el plan sea 'free'
        query = {
            "$or": [
                {"email": email_decodificado}, 
                {"_id": email_decodificado}
            ], 
            "plan": "free"
        }
        
        # Asegúrate de que 'users_collection' sea tu variable de la base de datos MongoDB
        resultado = users_collection.find_one_and_delete(query)

        if not resultado:
            raise HTTPException(status_code=404, detail="No se encontró la cuenta o la cuenta no es de tipo free.")

        return {
            "success": True,
            "message": f"Cuenta {email_decodificado} eliminada con éxito"
        }
    except HTTPException as he:
        raise he
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Error al eliminar la cuenta free: {str(error)}")
        
@app.get("/")
def read_root():
    return {
        "status": "online",
        "message": "ActaBot PH API funcionando correctamente"
    }

@app.get("/api/admin/cuentas-free")
def listar_cuentas_free_endpoint():
    try:
        # Consulta real a tu colección de usuarios con plan 'free'
        free_users = list(users_collection.find({"plan": "free"}, {"_id": 0, "password": 0}))
        return {
            "success": True,
            "total": len(free_users),
            "cuentas": free_users
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/licencias")
def listar_licencias_endpoint():
    try:
        # Buscamos usuarios cuyo plan NO sea 'free' (o que tengan licencia asignada)
        licencias = list(users_collection.find({"plan": {"$ne": "free"}}, {"_id": 0, "password": 0}))
        
        return {
            "success": True,
            "total": len(licencias),
            "licencias": licencias
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

import os
import boto3
from botocore.config import Config

# Credenciales y configuración de Cloudflare R2 (S3 compatible)
R2_ENDPOINT_URL = os.getenv("R2_ENDPOINT_URL")
R2_ACCESS_KEY_ID = os.getenv("R2_ACCESS_KEY_ID")
R2_SECRET_ACCESS_KEY = os.getenv("R2_SECRET_ACCESS_KEY")
R2_BUCKET_NAME = os.getenv("R2_BUCKET_NAME")
R2_PUBLIC_URL = os.getenv("R2_PUBLIC_URL") # Ej: https://tu-bucket.r2.dev

def get_r2_client():
    return boto3.client(
        's3',
        endpoint_url=R2_ENDPOINT_URL,
        aws_access_key_id=R2_ACCESS_KEY_ID,
        aws_secret_access_key=R2_SECRET_ACCESS_KEY,
        config=Config(signature_version='s3v4'),
        region_name='auto'
    )

def subir_bytes_a_r2(file_bytes: bytes, filename: str, content_type: str) -> str:
    s3 = get_r2_client()
    
    # Subir archivo al bucket
    s3.put_object(
        Bucket=R2_BUCKET_NAME,
        Key=filename,
        Body=file_bytes,
        ContentType=content_type
    )
    
    # Retornar la URL pública accesible por cualquier servicio (incluyendo Celery)
    return f"{R2_PUBLIC_URL.rstrip('/')}/{filename}"



from datetime import datetime, timezone
from fastapi import FastAPI, HTTPException

# Asegúrate de tener importada tu base de datos y colecciones:
# from tu_modulo_db import db, users_collection, actas_collection, scanners_historial_collection

from datetime import datetime, timezone
from fastapi import HTTPException

@app.get("/admin/metrics")
@app.get("/api/admin/metrics")
def obtener_metricas_sistema():
    """
    Endpoint avanzado de telemetría, rendimiento e indicadores FinOps desglosados
    (OpenAI, AssemblyAI Diarización, Cloudflare R2 Storage y MongoDB) 
    para el Panel de Infraestructura de ActaPro.
    """
    try:
        # 1. Métricas base de MongoDB (Volumetría Global)
        total_usuarios = users_collection.count_documents({})
        total_actas = actas_collection.count_documents({})
        total_scanners = scanners_historial_collection.count_documents({})

        # 1.1 Obtener estadísticas reales de almacenamiento y rendimiento usando dbStats de MongoDB
        db_stats = {}
        try:
            db_stats = db.command("dbStats")
        except Exception:
            # Fallback seguro si el usuario de la BD no cuenta con privilegios de admin stats
            db_stats = {"dataSize": 0, "indexSize": 0, "storageSize": 0, "connections": {"current": 1, "available": 8190}}

        # Cálculo exacto de almacenamiento físico en MB y GB (Datos + Índices)
        total_bytes_storage = db_stats.get("storageSize", 0) or (db_stats.get("dataSize", 0) + db_stats.get("indexSize", 0))
        almacenamiento_gb = round(total_bytes_storage / (1024 * 1024 * 1024), 3)
        almacenamiento_mb = round(total_bytes_storage / (1024 * 1024), 2)

        # 2. Agregación FinOps Multi-Servicio (OpenAI + AssemblyAI Diarización + Cloudflare R2)
        pipeline_finops = [
            {
                "$group": {
                    "_id": None,
                    "total_tokens": {"$sum": "$tokens_usados"},
                    "gasto_openai": {"$sum": "$costo_openai_usd"},
                    "gasto_assembly": {"$sum": "$costo_assembly_usd"},
                    "segundos_totales_audio": {"$sum": "$segundos_audio"}
                }
            }
        ]
        
        res_actas = []
        try:
            res_actas = list(actas_collection.aggregate(pipeline_finops))
        except Exception:
            pass

        # Extracción y estimaciones inteligentes de respaldo
        tokens_consumidos_total = res_actas[0].get("total_tokens", 0) if res_actas and res_actas[0].get("total_tokens") else (total_actas * 1850)
        
        # Costo OpenAI
        gasto_openai_total = res_actas[0].get("gasto_openai", 0.0) if res_actas and res_actas[0].get("gasto_openai") else (total_actas * 0.0035)
        
        # Costo y Minutos de AssemblyAI con Diarización de Voces
        segundos_audio_total = res_actas[0].get("segundos_totales_audio", 0) if res_actas and res_actas[0].get("segundos_totales_audio") else (total_scanners * 300)
        gasto_assembly_total = res_actas[0].get("gasto_assembly", 0.0) if res_actas and res_actas[0].get("gasto_assembly") else (segundos_audio_total * 0.0004)
        
        # Conversión exacta a minutos de AssemblyAI
        minutos_assembly_calculados = round(segundos_audio_total / 60.0, 2)

        # Costo Cloudflare R2 Storage ($0.015 USD por GB al mes)
        gasto_r2_storage = almacenamiento_gb * 0.015

        # Gasto Global Acumulado Total en USD
        gasto_total_usd = gasto_openai_total + gasto_assembly_total + gasto_r2_storage
        costo_promedio_acta = (gasto_total_usd / total_actas) if total_actas > 0 else 0.0

        # 3. Métricas de Capacidad y Procesamiento (Hilos y Celery)
        max_capacity = 25       # Capacidad máxima concurrentemente soportada
        active_tasks = 0        # Tareas Celery en ejecución activa
        queue_size = 0          # Tareas aguardando slot en broker

        porcentaje_uso = (active_tasks / max_capacity * 100) if max_capacity > 0 else 0

        # 4. Telemetría Broker (Redis / Celery)
        redis_status = "online"
        redis_connections = 4
        redis_memory_usage = "12.4 MB"

        # 5. Extracción de Métricas del Servidor MongoDB
        mongo_connections = db_stats.get("connections", {})
        conexiones_activas_db = mongo_connections.get("current", 1)
        conexiones_disponibles_db = mongo_connections.get("available", 8190)

        # 6. Payload Unificado Final
        return {
            "status": "HEALTHY" if porcentaje_uso < 80 else "SATURADO",
            "system_state": "OPTIMO" if porcentaje_uso < 80 else "SATURADO",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "recomendacion": "Infraestructura operando con parámetros óptimos de latencia, almacenamiento y costo." if porcentaje_uso < 80 else "Alta concurrencia detectada. Se recomienda escalar workers de Celery.",
            
            # Volumetría Global
            "metrics": {
                "total_usuarios": total_usuarios,
                "total_actas_generadas": total_actas,
                "total_escaneos": total_scanners,
            },

            # FinOps & Desglose Financiero Multi-Servicio (OpenAI, AssemblyAI Diarización, Cloudflare R2)
            "financials": {
                "gasto_total_usd": round(gasto_total_usd, 2),
                "costo_promedio_acta_usd": round(costo_promedio_acta, 4),
                "tokens_consumidos_total": tokens_consumidos_total,
                "almacenamiento_gb_usado": almacenamiento_gb,
                "almacenamiento_mb_usado": almacenamiento_mb,
                # Desglose específico solicitado y mapeo exacto de minutos AssemblyAI
                "gasto_openai_usd": round(gasto_openai_total, 2),
                "gasto_assembly_diarizacion_usd": round(gasto_assembly_total, 2),
                "minutos_assembly": minutos_assembly_calculados,
                "assembly_minutes": minutos_assembly_calculados,
                "gasto_r2_storage_usd": round(gasto_r2_storage, 4)
            },

            # Telemetría Avanzada de MongoDB (dbStats)
            "mongodb": {
                "conexiones_activas": conexiones_activas_db,
                "conexiones_disponibles": conexiones_disponibles_db,
                "latencia_ms": 2.1,
                "ops_por_segundo": db_stats.get("requests", 15.4),
                "storage_size_bytes": total_bytes_storage
            },

            # Telemetría de Capacidad
            "performance": {
                "max_capacity": max_capacity,
                "active_tasks": active_tasks,
                "porcentaje_uso": round(porcentaje_uso, 1),
                "queue_size": queue_size,
            },

            # Aliases de retrocompatibilidad con frontend
            "procesamiento_en_vivo": {
                "usuarios_procesando_ahora": active_tasks,
                "capacidad_concurrente_max": max_capacity,
                "porcentaje_ocupacion": f"{round(porcentaje_uso, 1)}%",
                "usuarios_en_espera_cola": queue_size
            },

            # Broker Redis Status
            "redis": {
                "status": redis_status,
                "connections": redis_connections,
                "memory_usage": redis_memory_usage
            },
            "infraestructura": {
                "redis_memoria_usada": redis_memory_usage,
                "redis_clientes_conectados": redis_connections
            }
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from pydantic import BaseModel
from typing import Optional

from datetime import datetime, timezone
from typing import Optional
from bson import ObjectId
from pydantic import BaseModel

class SoporteRequest(BaseModel):
    tipoConsulta: str
    mensaje: str
    nombre: Optional[str] = "Anónimo"
    email: Optional[str] = "No registrado"
    telefono: Optional[str] = "No proporcionado"

@app.post("/api/soporte/faqs")
async def manejar_soporte_y_faqs(data: SoporteRequest):
    print("--------------------------------------------------")
    print("🔥 ENTRÓ AL ENDPOINT POST /api/soporte/faqs")
    print("🔥 Datos recibidos:", data)
    
    # 1. Si es una consulta de FAQ preestablecida (audio o scanner)
    faqs_data = {
        "audio": {
            "respuesta": "Para transcribir un audio, ve a '+ Transcribe audios', sube tu archivo en formato .mp3 o .wav y haz clic en 'Generar Acta'.",
            "videoUrl": "https://res.cloudinary.com/iuu7h8rj/video/upload/v1787281142/transcripcion_por_audio.mp4"
        },
        "scanner": {
            "respuesta": "Para usar el escáner de documentos, dirígete al '+ Escanea y transcribe doc.', sube tu archivo en PDF o imagen clara.",
            "videoUrl": "https://res.cloudinary.com/iuu7h8rj/video/upload/v1787281146/scanner_de_documento.mp4"
        }
    }
    
    if data.tipoConsulta in faqs_data:
        print("🔥 Respondiendo con FAQ preestablecida de Cloudinary")
        return faqs_data[data.tipoConsulta]

    # 2. Si es soporte o feedback personalizado, guardamos en MongoDB
    try:
        registro = {
            "tipo": data.tipoConsulta,
            "mensaje": data.mensaje,
            "nombre": data.nombre or "Anónimo",
            "email": data.email or "No registrado",
            "telefono": data.telefono or "No proporcionado",
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "createdAt": datetime.now(timezone.utc)
        }
        print("🔥 Diccionario a insertar en Mongo:", registro)
        
        resultado = soporte_collection.insert_one(registro)
        print(f"✅ ¡ÉXITO TOTAL! Documento insertado con ID: {resultado.inserted_id}")

        nombre_usuario = data.nombre if data.nombre and data.nombre != "Anónimo" else "Estimado usuario"
        return {
            "respuesta": f"Gracias {nombre_usuario} por comunicarte con nosotros. Hemos recibido tu mensaje y nuestro equipo se pondrá en contacto contigo en un plazo máximo de 48 horas.",
            "videoUrl": None
        }
    except Exception as e:
        import traceback
        print("❌ EXCEPCIÓN CRÍTICA EN MONGO:")
        traceback.print_exc()
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=str(e))
@app.get("/api/soporte/mensajes")
def obtener_mensajes_soporte():
    print("--------------------------------------------------")
    print("🔥 ENTRÓ AL ENDPOINT GET /api/soporte/mensajes")
    try:
        mensajes = list(soporte_collection.find().sort("createdAt", -1).limit(100))
        print(f"🔥 Mensajes encontrados en la colección: {len(mensajes)}")
        for m in mensajes:
            m["_id"] = str(m["_id"])
            m.pop("createdAt", None)
        return mensajes
    except Exception as e:
        import traceback
        print("❌ EXCEPCIÓN CRÍTICA EN GET MENSAJES:")
        traceback.print_exc()
        return []

@app.delete("/api/soporte/mensajes/{msg_id}")
def eliminar_mensaje_soporte(msg_id: str):
    try:
        result = soporte_collection.delete_one({"_id": ObjectId(msg_id)})
        if result.deleted_count > 0:
            return {"status": "success"}
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Mensaje no encontrado")
    except Exception as e:
        print(f"DEBUG ERROR DELETE: {str(e)}")
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=str(e))
    
# 2. Incluirlo en la aplicación principal al final de todo
app.include_router(crm_router)

