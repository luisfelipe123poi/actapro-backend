import os
import io
import uuid
import base64
import hashlib
from datetime import datetime, timezone
from io import BytesIO

import boto3
from botocore.config import Config
import requests
import pymupdf as fitz
import pdfplumber
from docx import Document
from pymongo import MongoClient
import assemblyai as aai
import openai

from celery_app import celery_app

# ==========================================
# 1. CONFIGURACIÓN Y CREDENCIALES (ENV)
# ==========================================
R2_ENDPOINT_URL = os.getenv("R2_ENDPOINT_URL")
R2_ACCESS_KEY_ID = os.getenv("R2_ACCESS_KEY_ID")
R2_SECRET_ACCESS_KEY = os.getenv("R2_SECRET_ACCESS_KEY")
R2_BUCKET_NAME = os.getenv("R2_BUCKET_NAME", "archivos-temporales-actaprocore")
R2_PUBLIC_URL = os.getenv("R2_PUBLIC_URL", "https://cdn.actaprocore.com")

AAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")

aai.settings.api_key = AAI_API_KEY
openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
mongo_client = MongoClient(MONGO_URI)
db = mongo_client["actabot_db"]

# Colecciones de MongoDB
users_collection = db["users"]
actas_collection = db["actas_historial"]
transripciones_collection = db["transripciones_cache"]
scanners_historial_collection = db["scanners_historial"]

PROMPT_SISTEMA_ACTAS = """
Eres un Secretario Jurídico experto en Propiedad Horizontal en Colombia (Ley 675 de 2001). 
Tu objetivo es redactar un acta de asamblea formal, institucional, jurídicamente rigurosa y totalmente profesional a partir de la transcripción provista.

REGLAS DE FORMATO Y CONTENIDO OBLIGATORIAS:
1. NO UTILICES símbolos de almohadilla (#), asteriscos (*), guiones de lista markdown u otros caracteres de formato crudo. La redacción debe ser texto plano institucional, estructurado con títulos claros en mayúsculas sostenidas para cada sección.
2. En CADA punto tratado del orden del día, antes de pasar al siguiente punto, debes concluir obligatoriamente con dos apartados explícitos redactados de forma formal:
   - DECISIONES: (Detalla de forma precisa lo aprobado o resuelto en este punto).
   - PENDIENTES: (Detalla las tareas, responsables o acciones que quedaron abiertas en este punto, o indica formalmente si no aplica).
3. Al finalizar toda la reunión, antes de las firmas o el cierre, debes incluir obligatoriamente una sección final titulada exactamente:
   RESUMEN GENERAL DE DECISIONES Y PENDIENTES DE LA REUNION
   Donde consolides en forma de párrafos o listas institucionales limpias un resumen integrado de todas las decisiones tomadas y los pendientes asignados en la asamblea.
4. Mantén un tono neutro, impersonal, objetivo y estrictamente corporativo.
"""

def get_r2_client():
    return boto3.client(
        's3',
        endpoint_url=R2_ENDPOINT_URL,
        aws_access_key_id=R2_ACCESS_KEY_ID,
        aws_secret_access_key=R2_SECRET_ACCESS_KEY,
        config=Config(signature_version='s3v4'),
        region_name='auto'
    )


# ==========================================
# 2. TAREA: PROCESAR ASAMBLEA (AUDIO -> DOCX)
# ==========================================
@celery_app.task(
    bind=True,
    max_retries=5,
    default_retry_delay=15,
    autoretry_for=(requests.RequestException, openai.APIError, Exception)
)
def task_procesar_asamblea(self, temp_audio_path: str, email: str, instrucciones: str, nombre_personalizado: str, original_filename: str):
    try:
        self.update_state(state="PROCESSING", meta={"status": "Procesando audio e identificando oradores desde la nube..."})

        # Descarga temporal para hashing / caché
        response_audio = requests.get(temp_audio_path, timeout=60)
        if response_audio.status_code != 200:
            raise Exception(f"No se pudo descargar el archivo de audio desde la nube: {temp_audio_path}")
            
        content_bytes = response_audio.content
        file_hash = hashlib.sha256(content_bytes).hexdigest()
        cached = transripciones_collection.find_one({"file_hash": file_hash})

        duracion_segundos = 0

        if cached:
            texto_transcrito = cached["texto_transcrito"]
            duracion_segundos = cached.get("duracion_segundos", 300)
        else:
            config = aai.TranscriptionConfig(speaker_labels=True, language_code="es")
            transcriber = aai.Transcriber()
            
            transcript = transcriber.transcribe(temp_audio_path, config=config)

            if transcript.status == aai.TranscriptStatus.error:
                raise Exception(f"Error en AssemblyAI: {transcript.error}")

            # Calcular duración en segundos
            duracion_segundos = getattr(transcript, 'audio_duration', 300) or 300

            texto_transcrito = ""
            if transcript.utterances:
                for utterance in transcript.utterances:
                    texto_transcrito += f"[Persona {utterance.speaker}]: {utterance.text}\n"
            else:
                texto_transcrito = transcript.text

            transripciones_collection.insert_one({
                "file_hash": file_hash,
                "filename": original_filename,
                "texto_transcrito": texto_transcrito,
                "duracion_segundos": duracion_segundos,
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "createdAt": datetime.now(timezone.utc),
            })

        self.update_state(state="PROCESSING", meta={"status": "Generando el Acta con IA Jurídica..."})

        session_id = str(uuid.uuid4())
        nombre_archivo_acta = f"Acta_Asamblea_{session_id[:8]}.docx" if not nombre_personalizado else f"{nombre_personalizado.strip().replace(' ', '_')}.docx"

        prompt_final = PROMPT_SISTEMA_ACTAS
        if instrucciones:
            prompt_final += f"\n\nINSTRUCCIONES ADICIONALES DEL USUARIO:\n{instrucciones}"

        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt_final},
                {"role": "user", "content": f"Transcripción de la asamblea:\n\n{texto_transcrito}"},
            ],
            temperature=0.3,
        )
        acta_final = response.choices[0].message.content

        # Crear documento Word en memoria
        doc = Document()
        doc.add_heading("ACTA DE ASAMBLEA GENERAL DE COPROPIETARIOS", level=0).alignment = 1
        for linea in acta_final.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        docx_bytes = doc_io.read()

        # Subir .docx a Cloudflare R2
        r2_docx_key = f"actas_generadas/{session_id}_{nombre_archivo_acta}"
        s3 = get_r2_client()
        s3.put_object(
            Bucket=R2_BUCKET_NAME,
            Key=r2_docx_key,
            Body=docx_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        docx_url = f"{R2_PUBLIC_URL.rstrip('/')}/{r2_docx_key}"
        peso_archivo = f"{round(len(docx_bytes) / 1024, 1)} KB"

        # Cálculo de consumo de tiempo (Horas)
        duracion_horas = round(duracion_segundos / 3600.0, 2)
        if duracion_horas <= 0:
            duracion_horas = 0.01

        # Actualizar cuota de usuario en MongoDB
        if email:
            users_collection.update_one(
                {"email": email},
                {
                    "$inc": {
                        "horas_usadas_mes": duracion_horas,
                        "horas_restantes": -duracion_horas
                    }
                }
            )

        # Guardar historial de acta
        data_acta = {
            "email": email,
            "nombre_acta": nombre_archivo_acta,
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "peso": peso_archivo,
            "contenido": acta_final,
            "duracion_horas": duracion_horas,
            "file_url": docx_url
        }
        inserted = actas_collection.insert_one(data_acta)

        return {
            "status": "COMPLETED",
            "acta_id": str(inserted.inserted_id),
            "nombre_acta": nombre_archivo_acta,
            "file_url": docx_url
        }

    except Exception as exc:
        raise Exception(f"Fallo en la tarea de procesamiento: {str(exc)}")


# ==========================================
# 3. TAREA: ESCANEAR DOCUMENTO (OCR / HTML)
# ==========================================
@celery_app.task(
    bind=True,
    max_retries=5,
    default_retry_delay=10,
    autoretry_for=(openai.APIError, Exception)
)
def task_escanear_documento(self, file_bytes_b64: str, filename: str, email: str = None):
    try:
        self.update_state(state="PROCESSING", meta={"status": "Validando permisos y leyendo archivo..."})

        # Validar límite de tokens del usuario
        if email:
            usuario = users_collection.find_one({"email": email})
            if usuario:
                tokens_usados = usuario.get("tokens_usados", 0)
                limite_tokens = usuario.get("limite_tokens_mes", 0)
                if limite_tokens > 0 and tokens_usados >= limite_tokens:
                    raise Exception("Has alcanzado el límite de tokens mensuales de tu plan. Actualiza tu suscripción para continuar.")

        file_bytes = base64.b64decode(file_bytes_b64)
        filename_lower = filename.lower()
        
        texto_extraido = ""
        es_imagen = filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp"))
        
        if es_imagen:
            base64_image = base64.b64encode(file_bytes).decode("utf-8")
            contenido_usuario = [
                {
                    "type": "text",
                    "text": "Analyze this scanned document or image. Extract the information by structuring the visual design with corporate semantic HTML tags (use <h1>, <h2>, <p>, <table>, <thead>, <tbody>, <tr>, <th>, <td>). Apply Tailwind CSS classes to maintain a professional style (e.g., fonts, clean borders, spacing). DO NOT use Markdown, DO NOT use asterisks, DO NOT use code blocks of any kind. If there are data tables, create them completely with HTML tags. If you detect statistical charts or diagrams, represent them with a structured div with the class 'p-4 border-2 border-dashed border-slate-300 bg-slate-50 text-center text-slate-500 rounded-lg text-xs my-4' indicating the content of the chart."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}"
                    }
                }
            ]
        else:
            if filename_lower.endswith(".pdf"):
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                try:
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        texto_pagina = page.get_text()
                        
                        if texto_pagina.strip():
                            texto_extraido += f"\n--- Página {page_num + 1} ---\n" + texto_pagina
                finally:
                    doc.close()
                
                # Respaldo con pdfplumber si la lectura simple extrae poco texto
                if len(texto_extraido.strip()) < 50:
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for i, page in enumerate(pdf.pages):
                            t = page.extract_text()
                            if t:
                                texto_extraido += f"\n--- Página (Tablas) {i + 1} ---\n" + t

            elif filename_lower.endswith((".txt", ".doc", ".docx")):
                texto_extraido = file_bytes.decode("utf-8", errors="ignore")
            else:
                raise Exception("Formato de archivo no soportado. Sube un PDF, imagen o documento de texto.")

            if not texto_extraido.strip():
                raise Exception("El documento está vacío o no se pudo extraer texto legible.")

            contenido_usuario = f"""Analyze the following text extracted from the document. Your output must be EXCLUSIVELY corporate semantic HTML ready to render directly in a browser or web container.
- Replicate the original visual and section structure.
- Use <h1>, 2> for main and section titles.
- Use <p> for paragraphs with Tailwind classes (e.g., text-slate-900, text-xs, leading-relaxed).
- Use complete table tags (<table>, <thead>, <tbody>, <tr>, <th>, <td>) with borders and corporate classes if there is structured data.
- If you detect references to charts, schemes, or diagrams, create them as a visual block with a dotted border.
- FORBIDDEN to use Markdown, asterisks (*), markdown list hyphens (#), or wrap the result in markdown code quotes.

Extracted text:
{texto_extraido[:15000]}"""

        self.update_state(state="PROCESSING", meta={"status": "Procesando HTML estructurado con GPT-4o..."})

        response_openai = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You are a Documentation Engineer and Web Designer expert in corporate digitalization. 
Your sole mission is to transform the input document information into a pure, clean, and professional HTML code block integrated with Tailwind CSS classes.
STRICT RULES:
1. Return ONLY valid HTML code. Do not include prior explanations or text outside of the HTML.
2. NEVER use Markdown syntax (*, #, -, ```html). The result must be plain text containing exclusively HTML markup.
3. Structure data tables with <table>, <thead>, <tbody>, <tr>, <th>, and <td> applying clean classes (e.g., border border-slate-300 p-2).
4. Represent detected charts using a <div> container with dotted borders and professional design.
5. Maintain absolute fidelity to the original document structure."""
                },
                {
                    "role": "user",
                    "content": contenido_usuario
                }
            ],
            temperature=0.0
        )

        resultado_html = response_openai.choices[0].message.content.strip()

        # Descontar tokens utilizados
        tokens_consumidos = 0
        if email and hasattr(response_openai, "usage") and response_openai.usage:
            tokens_consumidos = response_openai.usage.total_tokens
            users_collection.update_one(
                {"email": email},
                {"$inc": {"tokens_usados": tokens_consumidos}}
            )

        # Limpieza defensiva de tags de Markdown
        if resultado_html.startswith("```html"):
            resultado_html = resultado_html[7:]
        if resultado_html.startswith("```"):
            resultado_html = resultado_html[3:]
        if resultado_html.endswith("```"):
            resultado_html = resultado_html[:-3]
        resultado_html = resultado_html.strip()

        # Guardar registro en historial de scanners
        scanner_id = None
        if email:
            nuevo_registro = {
                "email": email,
                "nombre": filename or "Documento Escaneado",
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "tokens": tokens_consumidos,
                "contenido": resultado_html
            }
            resultado_db = scanners_historial_collection.insert_one(nuevo_registro)
            scanner_id = str(resultado_db.inserted_id)

        return {
            "status": "COMPLETED",
            "transcripcion": resultado_html,
            "id": scanner_id
        }

    except Exception as e:
        raise Exception(f"Error procesando el archivo en Celery: {str(e)}")
